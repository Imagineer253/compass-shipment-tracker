from flask import Blueprint, render_template, request, redirect, url_for, flash, send_file, jsonify, current_app
from flask_login import login_required, current_user
from datetime import datetime
import os
from docx import Document
from docxtpl import DocxTemplate
from docx.shared import Cm, Pt
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.table import _Cell
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
import io
import json
from num2words import num2words
from .models import User, Role, Shipment, CombinedShipmentCounter, SigningAuthority, PackageQRCode, db
from werkzeug.security import generate_password_hash, check_password_hash
from .utils.helpers import generate_file_reference_number
from .services.qr_service import QRCodeService

main = Blueprint('main', __name__)

def prevent_table_page_break(table):
    """
    Prevent table from being split across pages by setting table properties
    """
    try:
        # Get table properties
        tblPr = table._element.xpath("./w:tblPr")[0]
        
        # Add table positioning properties to keep table together
        tblpPr = OxmlElement('w:tblpPr')
        
        # Set table to keep together on one page
        cantSplit = OxmlElement('w:cantSplit')
        cantSplit.set(qn('w:val'), '1')
        tblPr.append(cantSplit)
        
        # Set all rows to not break across pages
        for row in table.rows:
            trPr = row._element.get_or_add_trPr()
            cantSplit = OxmlElement('w:cantSplit')
            cantSplit.set(qn('w:val'), '1')
            trPr.append(cantSplit)
            
    except Exception as e:
        print(f"Warning: Could not set page break prevention: {e}")

def add_page_break_before_element(doc, target_element):
    """
    Add a page break before the target element if there might not be enough space
    """
    try:
        # Create a new paragraph before the target element
        new_paragraph = doc.add_paragraph()
        
        # Add a page break to the paragraph
        run = new_paragraph.add_run()
        run.add_break(WD_BREAK.PAGE)
        
        # Insert the new paragraph before the target element
        parent = target_element.getparent()
        parent.insert(list(parent).index(target_element), new_paragraph._element)
        
    except Exception as e:
        print(f"Warning: Could not add page break: {e}")

def get_package_type_display_name(package_type, form_data=None, package_num=None):
    """Convert package type code to display name, handling 'other' type with custom input"""
    package_type_mapping = {
        'cardboard_box': 'Cardboard Box',
        'plastic_crate': 'Plastic Crate',
        'metal_trunk': 'Metal Trunk',
        'zarges': 'Zarges',
        'pelican_case': 'Pelican Case',
        'other': 'Other',
        # Legacy support for old values
        'box': 'Cardboard Box',
        'carton': 'Plastic Crate',
        'crate': 'Metal Trunk'
    }
    
    # If package type is 'other' and we have form_data and package_num, get the custom type
    if package_type == 'other' and form_data and package_num:
        other_type = form_data.get(f'package_{package_num}_other_type', '')
        if other_type:
            return other_type
    
    return package_type_mapping.get(package_type, package_type.title())

def admin_required(f):
    """Decorator to require admin access"""
    def decorated_function(*args, **kwargs):
        if not current_user.is_authenticated or not current_user.is_admin():
            flash('Access denied. Admin privileges required.', 'error')
            return redirect(url_for('main.dashboard'))
        return f(*args, **kwargs)
    decorated_function.__name__ = f.__name__
    return decorated_function

@main.route('/shipment-type')
@login_required
def shipment_type_selection():
    return render_template('shipments/type_selection.html', user=current_user)

@main.route('/export-shipment')
@login_required
def export_shipment():
    # Get all active users for dropdown selection (both admin and regular users need this)
    all_users = User.query.filter_by(is_active=True).order_by(User.first_name, User.last_name).all()
    
    if current_user.is_admin():
        return render_template('shipments/admin_export_shipment.html', user=current_user, all_users=all_users)
    else:
        return render_template('shipments/export_shipment.html', user=current_user, all_users=all_users)

@main.route('/import-shipment')
@login_required
def import_shipment():
    # Get all active users for dropdown selection (both admin and regular users need this)
    all_users = User.query.filter_by(is_active=True).order_by(User.first_name, User.last_name).all()
    return render_template('shipments/import_shipment.html', user=current_user, all_users=all_users)

@main.route('/reimport-shipment')
@login_required
def reimport_shipment():
    # Get exported shipments for selection
    if current_user.is_admin():
        # Admin sees all exported shipments that have been delivered
        exported_shipments = Shipment.query.filter(
            Shipment.shipment_type == 'export',
            Shipment.status.in_(['Delivered', 'Document_Generated'])  # Only completed exports can be reimported
        ).order_by(Shipment.created_at.desc()).all()
    else:
        # Regular users see only their own exported shipments
        exported_shipments = Shipment.query.filter(
            Shipment.shipment_type == 'export',
            Shipment.created_by == current_user.id,
            Shipment.status.in_(['Delivered', 'Document_Generated'])  # Only completed exports can be reimported
        ).order_by(Shipment.created_at.desc()).all()
    
    return render_template('shipments/reimport_selection.html', 
                         user=current_user, 
                         exported_shipments=exported_shipments)

@main.route('/reimport-shipment/<int:shipment_id>')
@login_required
def reimport_form(shipment_id):
    # Get the selected shipment for reimport
    if current_user.is_admin():
        # Admin can reimport any exported shipment
        original_shipment = Shipment.query.filter(
            Shipment.id == shipment_id,
            Shipment.shipment_type == 'export',
            Shipment.status.in_(['Delivered', 'Document_Generated'])
        ).first_or_404()
    else:
        # Regular users can only reimport their own shipments
        original_shipment = Shipment.query.filter(
            Shipment.id == shipment_id,
            Shipment.shipment_type == 'export',
            Shipment.created_by == current_user.id,
            Shipment.status.in_(['Delivered', 'Document_Generated'])
        ).first_or_404()
    
    # Get all users for dropdown (needed for reimport form)
    all_users = User.query.filter_by(is_active=True).order_by(User.first_name, User.last_name).all()
    
    # Parse original shipment form data if available
    original_form_data = {}
    if original_shipment and original_shipment.form_data:
        try:
            import json
            original_form_data = json.loads(original_shipment.form_data)
        except (json.JSONDecodeError, TypeError):
            original_form_data = {}
    
    return render_template('shipments/reimport_shipment.html', 
                         user=current_user, 
                         original_shipment=original_shipment,
                         original_form_data=original_form_data,
                         all_users=all_users,
                         is_reimport=True)

@main.route('/cold-shipment')
@login_required
def cold_shipment():
    return render_template('shipments/cold_shipment.html', user=current_user)

def handle_table_placement(doc, form_data):
    """
    Creates a cargo table with precise width specifications.
    """
    try:
        # Find the placeholder paragraph
        target_paragraph = None
        for paragraph in doc.paragraphs:
            if '[INVOICE_TABLE]' in paragraph.text:
                target_paragraph = paragraph
                break
        
        if not target_paragraph:
            raise ValueError("Could not find [INVOICE_TABLE] placeholder in template")

        # Create the table with the Table Grid style
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        table.autofit = False  # Prevent autofit from overriding our widths
        
        # Define column widths in centimeters
        columns_width = {
            0: 2.0,   # Marks & no
            1: 3.0,   # No & kind of packages
            2: 7.34,  # Description of goods
            3: 2.5,   # Quantity
            4: 2.25,  # Rate USD
            5: 2.25   # Amount USD
        }
        
        # Set table width and alignment
        table.width = Cm(19.34)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER  # Center align the entire table
        
        # Remove top border from the entire table
        tblPr = table._element.xpath("./w:tblPr")[0]
        tblBorders = tblPr.find(qn('w:tblBorders'))
        if tblBorders is None:
            tblBorders = OxmlElement('w:tblBorders')
            tblPr.append(tblBorders)
        
        # Remove existing top border if any
        existing_top = tblBorders.find(qn('w:top'))
        if existing_top is not None:
            tblBorders.remove(existing_top)
        
        # Add top border set to none at table level
        top_border = OxmlElement('w:top')
        top_border.set(qn('w:val'), 'none')
        top_border.set(qn('w:sz'), '0')
        tblBorders.append(top_border)
        
        # Set up headers with proper formatting
        headers = ['Marks & no', 'No& kind of packages', 'Description of goods', 
                  'Quantity', 'Rate\nUSD', 'Amount\nUSD']
        
        for i, text in enumerate(headers):
            cell = table.cell(0, i)
            # Set column width
            cell.width = Cm(columns_width[i])
            
            paragraph = cell.paragraphs[0]
            paragraph.text = text
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Apply bold formatting to header text
            run = paragraph.runs[0]
            run.bold = True
            run.font.size = Pt(11)  # Set font size
            
            # Set vertical alignment to center
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            tcVAlign = OxmlElement('w:vAlign')
            tcVAlign.set(qn('w:val'), "center")
            tcPr.append(tcVAlign)
            
            # Remove top border from each header cell as well
            tcBorders = tcPr.find(qn('w:tcBorders'))
            if tcBorders is None:
                tcBorders = OxmlElement('w:tcBorders')
                tcPr.append(tcBorders)
            
            # Remove existing top border if any
            existing_top = tcBorders.find(qn('w:top'))
            if existing_top is not None:
                tcBorders.remove(existing_top)
            
            # Add top border set to none
            top_border = OxmlElement('w:top')
            top_border.set(qn('w:val'), 'none')
            top_border.set(qn('w:sz'), '0')
            tcBorders.append(top_border)

        # Add page break prevention to the table
        prevent_table_page_break(table)
        
        # Insert the table after the placeholder and remove the placeholder
        table._element = target_paragraph._element.addnext(table._element)
        target_paragraph._element.getparent().remove(target_paragraph._element)
        
        return table
        
    except Exception as e:
        print(f"Error in handle_table_placement: {str(e)}")
        raise

def handle_pl_table_placement(doc, form_data):
    """
    Creates a packing list table with precise width specifications.
    """
    try:
        # Find the placeholder paragraph
        target_paragraph = None
        for paragraph in doc.paragraphs:
            if '[PL_TABLE]' in paragraph.text:
                target_paragraph = paragraph
                break
        
        if not target_paragraph:
            raise ValueError("Could not find [PL_TABLE] placeholder in template")

        # Create the table with the Table Grid style
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        table.autofit = False  # Prevent autofit from overriding our widths
        
        # Define column widths in centimeters for PL table
        columns_width = {
            0: 2.0,   # Marks & no
            1: 3.0,   # No & kind of packages
            2: 7.34,  # Description of goods
            3: 2.0,   # Quantity
            4: 3.0,   # Dimensions
            5: 2.0    # Weight
        }
        
        # Set table width and alignment
        table.width = Cm(19.34)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER  # Center align the entire table
        
        # Remove top border from the entire table
        tblPr = table._element.xpath("./w:tblPr")[0]
        tblBorders = tblPr.find(qn('w:tblBorders'))
        if tblBorders is None:
            tblBorders = OxmlElement('w:tblBorders')
            tblPr.append(tblBorders)
        
        # Remove existing top border if any
        existing_top = tblBorders.find(qn('w:top'))
        if existing_top is not None:
            tblBorders.remove(existing_top)
        
        # Add top border set to none at table level
        top_border = OxmlElement('w:top')
        top_border.set(qn('w:val'), 'none')
        top_border.set(qn('w:sz'), '0')
        tblBorders.append(top_border)
        
        # Set up headers with proper formatting
        headers = ['Marks & no', 'No& kind of packages', 'Description of goods', 
                  'Quantity', 'Dimension\n(L x B x H)\n(cm)', 'Weight\n(kg)']
        
        for i, text in enumerate(headers):
            cell = table.cell(0, i)
            # Set column width
            cell.width = Cm(columns_width[i])
            
            paragraph = cell.paragraphs[0]
            paragraph.text = text
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Apply bold formatting to header text
            run = paragraph.runs[0]
            run.bold = True
            run.font.size = Pt(11)  # Set font size
            
            # Set vertical alignment to center
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            tcVAlign = OxmlElement('w:vAlign')
            tcVAlign.set(qn('w:val'), "center")
            tcPr.append(tcVAlign)
            
            # Remove top border from each header cell as well
            tcBorders = tcPr.find(qn('w:tcBorders'))
            if tcBorders is None:
                tcBorders = OxmlElement('w:tcBorders')
                tcPr.append(tcBorders)
            
            # Remove existing top border if any
            existing_top = tcBorders.find(qn('w:top'))
            if existing_top is not None:
                tcBorders.remove(existing_top)
            
            # Add top border set to none
            top_border = OxmlElement('w:top')
            top_border.set(qn('w:val'), 'none')
            top_border.set(qn('w:sz'), '0')
            tcBorders.append(top_border)

        # Add page break prevention to the table
        prevent_table_page_break(table)
        
        # Insert the table after the placeholder and remove the placeholder
        table._element = target_paragraph._element.addnext(table._element)
        target_paragraph._element.getparent().remove(target_paragraph._element)
        
        return table
        
    except Exception as e:
        print(f"Error in handle_pl_table_placement: {str(e)}")
        raise

def populate_table_data(table, form_data):
    """
    Populates the table with form data while maintaining precise formatting.
    """
    try:
        # Define column widths in centimeters
        columns_width = {
            0: 2.0,   # Marks & no
            1: 3.0,   # No & kind of packages
            2: 7.34,  # Description of goods
            3: 2.5,   # Quantity
            4: 2.25,  # Rate USD
            5: 2.25   # Amount USD
        }

        # Calculate the total width of columns 3-5 for the nested table
        nested_table_width = sum(columns_width[i] for i in range(3, 6))

        total_amount = 0
        total_packages = int(form_data.get('total_packages', 0))

        # Process each package (one row per package, combining all items)
        for package_num in range(1, total_packages + 1):
            items_count = int(form_data.get(f'package_{package_num}_items_count', 0))
            
            if items_count == 0:
                continue
                
            # Combine all items in this package
            item_descriptions = []
            total_package_quantity = 0
            total_package_value = 0
            
            for item_num in range(1, items_count + 1):
                prefix = f'package_{package_num}_item_{item_num}'
                
                # Get item details
                description = form_data.get(f'{prefix}_description', '')
                hsn_code = form_data.get(f'{prefix}_hsn_code', '')
                quantity = int(form_data.get(f'{prefix}_quantity', 0))
                unit_value = float(form_data.get(f'{prefix}_unit_value', 0))
                
                # Build item description
                if description:
                    item_desc = f"{description}"
                    if hsn_code:
                        item_desc += f" (HSN: {hsn_code})"
                    if quantity > 1:
                        item_desc += f" - {quantity} nos"
                    item_descriptions.append(item_desc)
                
                # Sum up totals for this package
                total_package_quantity += quantity
                total_package_value += (quantity * unit_value)
            
            # Skip if no valid items
            if not item_descriptions:
                continue
                
            # Add new row for this package
            row = table.add_row()
            row_cells = row.cells
            
            # Apply column widths to the new row
            for idx, width in columns_width.items():
                cell = row_cells[idx]
                cell.width = Cm(width)
            
            # Serial number
            row_cells[0].text = f"{package_num}."
            row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Box details - Include requester name for combined shipments
            package_type = form_data.get(f'package_{package_num}_type', '')
            package_type_display = get_package_type_display_name(package_type, form_data, package_num)
            
            # Get package owner information
            box_text = f"Box-{package_num} ({package_type_display})"
            
            # Check if this is a combined shipment with package ownership info
            package_belongs_to_id = form_data.get(f'package_{package_num}_belongs_to')
            attn_field_name = f'package_{package_num}_attn'
            
            # Collect ownership info to add compactly
            ownership_info = []
            
            if package_belongs_to_id:
                # Combined shipment - get user by ID for package ownership
                try:
                    from .models import User
                    owner_user = User.query.get(int(package_belongs_to_id))
                    if owner_user:
                        owner_name = f"{owner_user.first_name} {owner_user.last_name}"
                        owner_unique_id = owner_user.unique_id
                        ownership_info.append(f"Owner: {owner_name} ({owner_unique_id})")
                except (ValueError, TypeError):
                    pass
                    
            # Check if there's attention field (requester) for items
            if attn_field_name in form_data and form_data[attn_field_name]:
                # Get requester unique ID if available
                requester_name = form_data[attn_field_name]
                try:
                    from .models import User
                    # Try to find user by full name to get their unique ID
                    name_parts = requester_name.split()
                    if len(name_parts) >= 2:
                        first_name = name_parts[0]
                        last_name = ' '.join(name_parts[1:])  # Handle multi-word last names
                        requester_user = User.query.filter_by(first_name=first_name, last_name=last_name).first()
                        if requester_user and requester_user.unique_id:
                            ownership_info.append(f"Req: {requester_name} ({requester_user.unique_id})")
                        else:
                            ownership_info.append(f"Req: {requester_name}")
                    else:
                        ownership_info.append(f"Req: {requester_name}")
                except Exception:
                    ownership_info.append(f"Req: {requester_name}")
            
            # Format the final box text more compactly
            if ownership_info:
                # Put ownership info on new line but in a more compact format
                box_text += f"\n{' | '.join(ownership_info)}"
            
            row_cells[1].text = box_text
            row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Combined description of all items in package
            combined_description = ", ".join(item_descriptions)
            row_cells[2].text = combined_description
            row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Total quantity for package
            row_cells[3].text = f"{total_package_quantity} SET"
            row_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Average rate for package (total value / total quantity)
            avg_rate = total_package_value / total_package_quantity if total_package_quantity > 0 else 0
            row_cells[4].text = f"{avg_rate:.0f}"
            row_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # Total amount for package
            row_cells[5].text = f"{total_package_value:.0f}"
            row_cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            total_amount += total_package_value

        # Convert number to words and capitalize first letter of each word, and "And"
        amount_words = num2words(int(total_amount), to='cardinal').split()
        capitalized_words = []
        for word in amount_words:
            if word.lower() == 'and':
                capitalized_words.append('And')
            else:
                capitalized_words.append(word.capitalize())
        amount_in_words = " ".join(capitalized_words) + " Only"

        # Create the declarations text with proper formatting
        declarations_text = (
            f"Amount Chargeable (in words): USD {amount_in_words}\n\n"
            "THE VALUE DECLARED ABOVE IS FOR CUSTOMS PURPOSE ONLY.\n\n"
            "GOODS ARE FOR RESEARCH AND DEVELOPMENT.\n\n"
            "Declaration:\n"
            "We declare that the invoice shows the actual price of goods\n"
            "Described and that all particulars are true and correct."
        )

        # Add final row for declarations
        final_row = table.add_row()
        final_cells = final_row.cells
        
        # Apply column widths
        for idx, width in columns_width.items():
            cell = final_cells[idx]
            cell.width = Cm(width)

        # Set up declarations section (columns 0-2)
        final_cells[0].text = declarations_text
        for i in range(1, 3):
            final_cells[0].merge(final_cells[i])
        final_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Merge the cells for the right side first
        final_cells[3].merge(final_cells[4])
        final_cells[3].merge(final_cells[5])

        # Get the exact width of the merged cell in twips (1/20th of a point)
        merged_cell_width = int(nested_table_width * 1440)  # Convert cm to twips
        
        # Create a 2x1 table for the right side (TOTAL and Signature)
        right_table = Document().add_table(rows=2, cols=2)
        right_table.style = None  # Remove any default styling that adds borders
        right_table.autofit = False

        # Set table properties to match container cell exactly and center align the table
        tblPr = right_table._element.xpath("./w:tblPr")[0]
        
        # Completely remove all table-level borders
        tblBorders = OxmlElement('w:tblBorders')
        for border in ['top', 'bottom', 'left', 'right', 'insideH', 'insideV']:
            border_elem = OxmlElement(f'w:tcBorders')
            border_elem.set(qn('w:val'), 'none')
            border_elem.set(qn('w:sz'), '0')
            tblBorders.append(border_elem)
        tblPr.append(tblBorders)
        
        # Convert centimeters to twips (1cm = 567 twips)
        total_width = int(7.0 * 567)  # Total table width of 7cm
        tblW = parse_xml(f'<w:tblW xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:w="{total_width}" w:type="dxa"/>')
        jc = parse_xml('<w:jc xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="center"/>')
        tblPr.append(tblW)
        tblPr.append(jc)
        
        # Set exact row heights
        total_row = right_table.rows[0]
        signature_row = right_table.rows[1]
        
        # Set TOTAL row height to 0.8cm
        total_row.height_rule = WD_ROW_HEIGHT.EXACTLY
        total_row.height = Cm(0.8)
        
        # Set signature row height to (3.64 - 0.8) = 2.84cm
        signature_row.height_rule = WD_ROW_HEIGHT.EXACTLY
        signature_row.height = Cm(2.84)

        # Set up TOTAL row with exact column widths
        total_label_cell = right_table.rows[0].cells[0]
        total_amount_cell = right_table.rows[0].cells[1]
        
        # Convert cell widths to twips
        total_label_width = int(2.5 * 567)  # 2.5cm for TOTAL label
        total_amount_width = int(4.5 * 567)  # 4.5cm for amount
        
        # Set exact widths for both cells
        total_label_cell._tc.get_or_add_tcPr().append(
            parse_xml(f'<w:tcW xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:w="{total_label_width}" w:type="dxa"/>')
        )
        total_amount_cell._tc.get_or_add_tcPr().append(
            parse_xml(f'<w:tcW xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:w="{total_amount_width}" w:type="dxa"/>')
        )
        
        # Format TOTAL label
        total_label_paragraph = total_label_cell.paragraphs[0]
        total_label_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = total_label_paragraph.add_run("TOTAL")
        run.font.bold = True
        
        # Format amount
        amount_paragraph = total_amount_cell.paragraphs[0]
        amount_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = amount_paragraph.add_run(f"{total_amount:.0f}")
        run.font.bold = True

        # Set up Signature row
        sig_cell = right_table.rows[1].cells[0]
        sig_cell.merge(right_table.rows[1].cells[1])
        
        # Set width for merged signature cell (7cm)
        sig_cell_width = int(7.0 * 567)  # 7cm
        sig_cell._tc.get_or_add_tcPr().append(
            parse_xml(f'<w:tcW xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:w="{sig_cell_width}" w:type="dxa"/>')
        )
        
        sig_paragraph = sig_cell.paragraphs[0]
        sig_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        sig_paragraph.add_run("Signature & date")
        
        # Add extra paragraphs for signature space
        for _ in range(2):
            sig_cell.add_paragraph()

        # Remove any existing content from container cell
        final_cells[3]._element.clear_content()
        
        # Convert the nested table to XML and insert it
        nested_table_xml = right_table._element
        final_cells[3]._tc.append(nested_table_xml)

        # Remove all borders from the nested table properly
        def remove_cell_borders(cell):
            """Remove all borders from a cell by setting them to 'none'"""
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            
            # Set all borders to none
            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border_elem = OxmlElement(f'w:{border_name}')
                border_elem.set(qn('w:val'), 'none')
                tcBorders.append(border_elem)
            
            tcPr.append(tcBorders)

        # Remove borders from all cells in the nested table except TOTAL row bottom border
        for i, row in enumerate(right_table.rows):
            for cell in row.cells:
                if i == 0:  # TOTAL row - remove all borders first, then add bottom border
                    remove_cell_borders(cell)
                else:  # Signature row - remove all borders
                    remove_cell_borders(cell)

        # Add bottom border under TOTAL row using Word's built-in properties
        def set_bottom_border(cell):
            """Set bottom border for a cell using Word's built-in properties"""
            tcPr = cell._tc.get_or_add_tcPr()
            # Find existing tcBorders or create new one
            tcBorders = tcPr.find(qn('w:tcBorders'))
            if tcBorders is None:
                tcBorders = OxmlElement('w:tcBorders')
                tcPr.append(tcBorders)
            
            # Remove existing bottom border if any
            existing_bottom = tcBorders.find(qn('w:bottom'))
            if existing_bottom is not None:
                tcBorders.remove(existing_bottom)
            
            # Add new bottom border
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '4')
            tcBorders.append(bottom)

        # Apply bottom border to both TOTAL cells
        set_bottom_border(total_label_cell)
        set_bottom_border(total_amount_cell)

        return total_amount

    except Exception as e:
        print(f"Error in populate_table_data: {str(e)}")
        raise

def populate_pl_table_data(table, form_data):
    """
    Populates the packing list table with form data while maintaining precise formatting.
    """
    try:
        # Define column widths in centimeters for PL table
        columns_width = {
            0: 2.0,   # Marks & no
            1: 3.0,   # No & kind of packages
            2: 7.34,  # Description of goods
            3: 2.0,   # Quantity
            4: 3.0,   # Dimensions
            5: 2.0    # Weight
        }

        total_packages = int(form_data.get('total_packages', 0))
        
        # Process each package (one row per package, combining all items)
        for package_num in range(1, total_packages + 1):
            items_count = int(form_data.get(f'package_{package_num}_items_count', 0))
            
            if items_count == 0:
                continue
                
            # Combine all items in this package
            item_descriptions = []
            total_package_quantity = 0
            total_package_weight = 0
            
            for item_num in range(1, items_count + 1):
                prefix = f'package_{package_num}_item_{item_num}'
                
                # Get item details
                description = form_data.get(f'{prefix}_description', '')
                hsn_code = form_data.get(f'{prefix}_hsn_code', '')
                quantity = int(form_data.get(f'{prefix}_quantity', 0))
                net_weight = float(form_data.get(f'{prefix}_net_weight', 0))
                
                # Build item description
                if description:
                    item_desc = f"{description}"
                    if hsn_code:
                        item_desc += f" (HSN: {hsn_code})"
                    if quantity > 1:
                        item_desc += f" - {quantity} nos"
                    item_descriptions.append(item_desc)
                
                # Sum up totals for this package
                total_package_quantity += quantity
                total_package_weight += (net_weight * quantity)  # Total weight = unit weight * quantity
            
            # Skip if no valid items
            if not item_descriptions:
                continue
                
            # Add new row for this package
            row = table.add_row()
            row_cells = row.cells
            
            # Apply column widths to the new row
            for idx, width in columns_width.items():
                cell = row_cells[idx]
                cell.width = Cm(width)
            
            # Serial number
            row_cells[0].text = f"{package_num}."
            row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Box details
            package_type = form_data.get(f'package_{package_num}_type', '')
            package_type_display = get_package_type_display_name(package_type, form_data, package_num)
            
            # Get package owner information for packing list
            box_text = f"Box-{package_num} ({package_type_display})"
            
            # Check if this is a combined shipment with package ownership info
            package_belongs_to_id = form_data.get(f'package_{package_num}_belongs_to')
            attn_field_name = f'package_{package_num}_attn'
            
            # Collect ownership info to add compactly
            ownership_info = []
            
            if package_belongs_to_id:
                # Combined shipment - get user by ID for package ownership
                try:
                    from .models import User
                    owner_user = User.query.get(int(package_belongs_to_id))
                    if owner_user:
                        owner_name = f"{owner_user.first_name} {owner_user.last_name}"
                        owner_unique_id = owner_user.unique_id
                        ownership_info.append(f"Owner: {owner_name} ({owner_unique_id})")
                except (ValueError, TypeError):
                    pass
                    
            # Check if there's attention field (requester) for items
            if attn_field_name in form_data and form_data[attn_field_name]:
                # Get requester unique ID if available
                requester_name = form_data[attn_field_name]
                try:
                    from .models import User
                    # Try to find user by full name to get their unique ID
                    name_parts = requester_name.split()
                    if len(name_parts) >= 2:
                        first_name = name_parts[0]
                        last_name = ' '.join(name_parts[1:])  # Handle multi-word last names
                        requester_user = User.query.filter_by(first_name=first_name, last_name=last_name).first()
                        if requester_user and requester_user.unique_id:
                            ownership_info.append(f"Req: {requester_name} ({requester_user.unique_id})")
                        else:
                            ownership_info.append(f"Req: {requester_name}")
                    else:
                        ownership_info.append(f"Req: {requester_name}")
                except Exception:
                    ownership_info.append(f"Req: {requester_name}")
            
            # Format the final box text more compactly
            if ownership_info:
                # Put ownership info on new line but in a more compact format
                box_text += f"\n{' | '.join(ownership_info)}"
            
            row_cells[1].text = box_text
            row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Combined description of all items in package
            combined_description = ", ".join(item_descriptions)
            row_cells[2].text = combined_description
            row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Total quantity for package
            row_cells[3].text = f"{total_package_quantity}"
            row_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Dimensions
            length = form_data.get(f'package_{package_num}_length', '')
            width = form_data.get(f'package_{package_num}_width', '')
            height = form_data.get(f'package_{package_num}_height', '')
            row_cells[4].text = f"{length} x {width} x {height}"
            row_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Total weight for package
            row_cells[5].text = f"{total_package_weight:.2f}"
            row_cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Add final row for declaration and signature
        final_row = table.add_row()
        final_cells = final_row.cells
        
        # Apply column widths
        for idx, width in columns_width.items():
            cell = final_cells[idx]
            cell.width = Cm(width)

        # Merge cells for declaration (left side)
        final_cells[0].merge(final_cells[1])
        final_cells[0].merge(final_cells[2])
        
        # Merge cells for signature (right side)
        final_cells[3].merge(final_cells[4])
        final_cells[3].merge(final_cells[5])
        
        # Add declaration text on the left
        declaration_text = (
            "THE ABOVE MENTIONED GOODS ARE FOR RESEARCH AND DEVELOPMENT.\n\n"
            "Declaration:\n"
            "We declare that all particulars given above are true and correct."
        )
        final_cells[0].text = declaration_text
        final_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Add signature text on the right
        final_cells[3].text = "Signature & date"
        final_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Set row height for signature space
        final_row.height = Cm(3.0)
        final_row.height_rule = WD_ROW_HEIGHT.EXACTLY
        
        return None
        
    except Exception as e:
        print(f"Error in populate_pl_table_data: {str(e)}")
        raise

def handle_shipper_table_placement(doc, form_data):
    """
    Finds the [Shipper_table] placeholder and replaces it with a table structure.
    """
    try:
        # Search for the placeholder in all paragraphs
        target_paragraph = None
        for paragraph in doc.paragraphs:
            if '[Shipper_table]' in paragraph.text:
                target_paragraph = paragraph
                break
        
        if not target_paragraph:
            print("Shipper_table placeholder not found in document")
            return None
        
        # Get the parent element (should be the document body)
        parent_element = target_paragraph._element.getparent()
        
        # Create a new table with 3 columns
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        table.autofit = False
        
        # Set column widths - adjust as needed
        col_widths = [Cm(4.0), Cm(10.0), Cm(3.0)]  # Column widths in cm
        for i, width in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = width
        
        # Set up header row
        header_cells = table.rows[0].cells
        
        # Column 1: "Marks and Number of Packages"
        header_cells[0].text = "Marks and Number of Packages"
        header_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_cells[0].paragraphs[0].runs[0].font.bold = True
        
        # Column 2: "Proper Description of Goods"
        header_cells[1].text = "Proper Description of Goods\n(Trade names not permitted)\nspecify each article separately"
        header_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_cells[1].paragraphs[0].runs[0].font.bold = True
        
        # Column 3: "Net Quantity"
        header_cells[2].text = "Net Quantity"
        header_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_cells[2].paragraphs[0].runs[0].font.bold = True
        
        # Add page break prevention to the table
        prevent_table_page_break(table)
        
        # Insert the table before the placeholder paragraph
        parent_element.insert(list(parent_element).index(target_paragraph._element), table._element)
        
        # Remove the placeholder paragraph
        parent_element.remove(target_paragraph._element)
        
        return table
        
    except Exception as e:
        print(f"Error handling shipper table placement: {e}")
        return None

def populate_shipper_table_data(table, form_data):
    """
    Populates the shipper table with package and item data.
    """
    try:
        if not table:
            return
            
        total_packages = int(form_data.get('total_packages', 0))
        
        # Process each package
        for package_num in range(1, total_packages + 1):
            items_count = int(form_data.get(f'package_{package_num}_items_count', 0))
            
            # Package details
            package_type = form_data.get(f'package_{package_num}_type', 'box')
            package_type_display = get_package_type_display_name(package_type, form_data, package_num)
            
            # Package dimensions
            length = form_data.get(f'package_{package_num}_length', '')
            width = form_data.get(f'package_{package_num}_width', '')
            height = form_data.get(f'package_{package_num}_height', '')
            dimensions = f"({length} x {width} x {height} cm)" if all([length, width, height]) else ""
            
            # Get package owner name (without label)
            owner_name = ""
            package_belongs_to_id = form_data.get(f'package_{package_num}_belongs_to')
            if package_belongs_to_id:
                try:
                    from .models import User
                    owner_user = User.query.get(int(package_belongs_to_id))
                    if owner_user:
                        owner_name = f" {owner_user.first_name} {owner_user.last_name}"
                except (ValueError, TypeError):
                    pass
            
            # Check for requester name (combined shipment)
            requester_info = ""
            attn_field_name = f'package_{package_num}_attn'
            if attn_field_name in form_data and form_data[attn_field_name]:
                requester_info = f" {form_data[attn_field_name]} – "
            
            # Create description content
            description_parts = []
            description_parts.append(f"Box No.- {package_num:02d} ({package_type_display}){requester_info}{dimensions}{owner_name}")
            
            # Add all items in this package
            for item_num in range(1, items_count + 1):
                prefix = f'package_{package_num}_item_{item_num}'
                description = form_data.get(f'{prefix}_description', '')
                hsn_code = form_data.get(f'{prefix}_hsn_code', '')
                quantity = form_data.get(f'{prefix}_quantity', '0')
                
                if description:
                    item_text = f"{description}"
                    if hsn_code:
                        item_text += f" (HSN: {hsn_code})"
                    if quantity and quantity != '0':
                        item_text += f" – {quantity} nos"
                    description_parts.append(item_text)
            
            # Add row for this package
            row = table.add_row()
            row_cells = row.cells
            
            # Set column widths for the new row
            col_widths = [Cm(4.0), Cm(10.0), Cm(3.0)]
            for i, width in enumerate(col_widths):
                row_cells[i].width = width
            
            # Column 1: "As Address"
            row_cells[0].text = "As Address"
            row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Column 2: Package and items description
            description_text = ", ".join(description_parts)
            row_cells[1].text = description_text
            row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Column 3: "1 set"
            row_cells[2].text = "1 set"
            row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        return table
        
    except Exception as e:
        print(f"Error populating shipper table: {e}")
        return table

@main.route('/submit-shipment', methods=['POST'])
@login_required
def submit_shipment():
    # Extract form data
    data = {
        'shipment_type': request.form.get('shipment_type'),
        'invoice_date': request.form.get('invoice_date'),
        'requester_name': request.form.get('requester_name'),
        'expedition_year': request.form.get('expedition_year'),
        'expedition_month': request.form.get('expedition_month'),
        'return_type': request.form.get('return_type'),
        'import_purpose': request.form.get('import_purpose'),
        'exporter': request.form.get('exporter'),
        'consignee': request.form.get('consignee'),
        'total_packages': request.form.get('total_packages'),
        'port_of_loading': request.form.get('port_of_loading'),
        'port_of_discharge': request.form.get('port_of_discharge'),
        'country_of_origin': request.form.get('country_of_origin'),
        'country_of_final_destination': request.form.get('country_of_final_destination'),
        'terms_of_delivery': request.form.get('terms_of_delivery'),
        'pre_carriage_by': request.form.get('pre_carriage_by'),
        'place_of_receipt': request.form.get('place_of_receipt'),
        'vessel_flight_no': request.form.get('vessel_flight_no'),
        'marks_numbers': request.form.get('marks_numbers'),
        'nature_of_goods': request.form.get('nature_of_goods'),
        'weight': request.form.get('weight'),
        'measurement': request.form.get('measurement'),
        'coldchain_required': request.form.get('coldchain_required') == 'yes'
    }

    # Generate invoice number based on shipment type with user's unique ID and serial number
    year = data['expedition_year']
    month = data['expedition_month']
    
    # Generate next serial number
    from compass.models import ShipmentSerialCounter, User
    serial_number = ShipmentSerialCounter.get_next_serial()
    
    # For admin users creating export shipments, check for combined export
    if current_user.is_admin() and data['shipment_type'] == 'export':
        # Get selected requester user ID
        requester_user_id = request.form.get('requester_user_id')
        
        # Extract package user assignments from form data
        package_user_assignments = {}
        total_packages = int(request.form.get('total_packages', 0))
        
        for i in range(1, total_packages + 1):
            user_id = request.form.get(f'package_{i}_belongs_to')
            if user_id:
                package_user_assignments[str(i)] = int(user_id)
        
        # Extract unique user IDs in order of first package assignment
        unique_user_ids = []
        seen_users = set()
        
        for package_num in sorted(package_user_assignments.keys(), key=int):
            user_id = package_user_assignments[package_num]
            if user_id and user_id not in seen_users:
                user = User.query.get(user_id)
                if user and user.unique_id:
                    unique_user_ids.append(user.unique_id)
                    seen_users.add(user_id)
        
        # Generate combined or single export invoice
        if len(unique_user_ids) > 1:
            # Combined export
            unique_ids_str = "/".join(unique_user_ids)
            invoice_number = f"NCPOR/ARC/{year}/{month}/EXP/{data['return_type']}/CMB/{unique_ids_str}/{serial_number}"
        elif len(unique_user_ids) == 1:
            # Single user export from package assignments
            invoice_number = f"NCPOR/ARC/{year}/{month}/EXP/{data['return_type']}/{unique_user_ids[0]}/{serial_number}"
        else:
            # Use selected requester's unique ID, fallback to current user
            if requester_user_id:
                selected_user = User.query.get(requester_user_id)
                user_unique_id = selected_user.unique_id if selected_user and selected_user.unique_id else 'XXXXXX'
            else:
                user_unique_id = current_user.unique_id if hasattr(current_user, 'unique_id') and current_user.unique_id else 'XXXXXX'
            invoice_number = f"NCPOR/ARC/{year}/{month}/EXP/{data['return_type']}/{user_unique_id}/{serial_number}"
    else:
        # Regular invoice generation for non-admin or non-export types
        user_unique_id = current_user.unique_id if hasattr(current_user, 'unique_id') and current_user.unique_id else 'XXXXXX'
        
        if data['shipment_type'] == 'export':
            invoice_number = f"NCPOR/ARC/{year}/{month}/EXP/{data['return_type']}/{user_unique_id}/{serial_number}"
        elif data['shipment_type'] == 'import':
            # Use new NCPOR/ARC pattern for import shipments with admin ID and requester ID
            admin_id = current_user.unique_id if current_user.unique_id else 'ADMIN'
            requester_unique_id = request.form.get('requester_unique_id', user_unique_id)
            invoice_number = f"NCPOR/ARC/{year}/{month}/SAM/RT/{admin_id}/{requester_unique_id}/{serial_number}"
        elif data['shipment_type'] == 'reimport':
            invoice_number = f"NCPOR/REIMP/{year}/{month}/{data['return_type']}/{user_unique_id}/{serial_number}"
        elif data['shipment_type'] == 'cold':
            invoice_number = f"NCPOR/COLD/{year}/{month}/{user_unique_id}/{serial_number}"
        else:
            invoice_number = f"NCPOR/UNKNOWN/{year}/{month}/{data['return_type']}/{user_unique_id}/{serial_number}"
    
    try:
        # Get form data
        form_data = request.form.to_dict()
        
        # Validate package limit - different limits for regular vs combined exports
        total_packages = int(form_data.get('total_packages', 0))
        export_type = form_data.get('export_type', 'regular')
        is_combined_export = (export_type == 'combined' and data['shipment_type'] == 'export' and current_user.is_admin())
        max_packages = 20 if is_combined_export else 10
        
        if total_packages > max_packages:
            export_type_text = 'combined export' if is_combined_export else 'shipment'
            flash(f'Error: Maximum {max_packages} packages allowed per {export_type_text}.', 'error')
            # Redirect back to the appropriate shipment form based on type
            if data['shipment_type'] == 'export':
                return redirect(url_for('main.export_shipment'))
            elif data['shipment_type'] == 'import':
                return redirect(url_for('main.import_shipment'))
            elif data['shipment_type'] == 'reimport':
                return redirect(url_for('main.reimport_shipment'))
            elif data['shipment_type'] == 'cold':
                return redirect(url_for('main.cold_shipment'))
            return redirect(url_for('main.shipment_type_selection'))
        
        # Ensure batch number is always uppercase
        batch_number = form_data.get('batch_number', '').upper()
        form_data['batch_number'] = batch_number
        
        # Create shipment record in database
        shipment = Shipment(
            invoice_number=invoice_number,
            serial_number=serial_number,
            shipment_type=data['shipment_type'],
            created_by=current_user.id,
            requester_name=form_data.get('requester_name', ''),
            expedition_year=form_data.get('expedition_year', ''),
            batch_number=form_data.get('batch_number', ''),
            destination_country=form_data.get('destination_country', 'NORWAY') if data['shipment_type'] == 'export' else 'INDIA',
            total_packages=int(form_data.get('total_packages', 0)),
            form_data=json.dumps(form_data),
            status='Submitted'
        )
        
        db.session.add(shipment)
        db.session.commit()
        
        # Generate QR codes for each package after shipment is created
        try:
            qr_service = QRCodeService()
            base_url = request.url_root.rstrip('/')  # Get the base URL without trailing slash
            
            total_packages = int(form_data.get('total_packages', 0))
            
            for package_num in range(1, total_packages + 1):
                # Extract package information from form data
                package_data = {
                    'type': form_data.get(f'package_{package_num}_type'),
                    'description': form_data.get(f'package_{package_num}_description'),
                    'weight': form_data.get(f'package_{package_num}_weight'),
                    'dimensions': form_data.get(f'package_{package_num}_dimensions'),
                    'attention_person_id': form_data.get(f'package_{package_num}_belongs_to')
                }
                
                # Generate QR code for this package
                package_qr = qr_service.generate_package_qr_code(
                    shipment=shipment,
                    package_number=package_num,
                    package_data=package_data,
                    base_url=base_url
                )
                
                if not package_qr:
                    current_app.logger.warning(f"Failed to generate QR code for package {package_num} in shipment {shipment.id}")
            
            current_app.logger.info(f"Generated QR codes for {total_packages} packages in shipment {shipment.id}")
            
        except Exception as e:
            current_app.logger.error(f"Error generating QR codes for shipment {shipment.id}: {str(e)}")
            # Don't fail the shipment creation if QR generation fails
        
        # All users (including admins) submit for review - no immediate document generation
        flash(f'Shipment {invoice_number} submitted successfully! Admin will review and acknowledge receipt.', 'success')
        return redirect(url_for('main.dashboard'))
        
    except Exception as e:
        # Update shipment status to failed if it was created
        try:
            if 'shipment' in locals():
                shipment.status = 'Failed'
                db.session.commit()
        except:
            pass
            
        flash(f'Error submitting shipment: {str(e)}', 'error')
        # Redirect back to the appropriate shipment form based on type
        if data['shipment_type'] == 'export':
            return redirect(url_for('main.export_shipment'))
        elif data['shipment_type'] == 'import':
            return redirect(url_for('main.import_shipment'))
        elif data['shipment_type'] == 'reimport':
            return redirect(url_for('main.reimport_shipment'))
        elif data['shipment_type'] == 'cold':
            return redirect(url_for('main.cold_shipment'))
        return redirect(url_for('main.shipment_type_selection'))

def get_aggregated_sample_types(form_data):
    """
    Extract and aggregate sample types from form data.
    Returns a formatted string of unique sample types.
    """
    sample_types = set()
    
    # Extract sample types from all items
    for key, value in form_data.items():
        if key.endswith('_sample_type') and value:
            if value == 'other':
                # Look for the corresponding other_sample_type field
                prefix = key.replace('_sample_type', '')
                other_key = f"{prefix}_other_sample_type"
                other_value = form_data.get(other_key, '').strip()
                if other_value:
                    sample_types.add(other_value.lower())
            else:
                # Map values to display names
                type_mapping = {
                    'water': 'water',
                    'sediment': 'sediment', 
                    'glass_fiber': 'glass fiber'
                }
                display_name = type_mapping.get(value, value)
                sample_types.add(display_name)
    
    # Convert to sorted list for consistent output
    sorted_types = sorted(list(sample_types))
    
    if not sorted_types:
        return "samples"
    elif len(sorted_types) == 1:
        return f"{sorted_types[0]} samples"
    else:
        # Join with "and" for the last item
        if len(sorted_types) == 2:
            return f"{sorted_types[0]} and {sorted_types[1]} samples"
        else:
            return f"{', '.join(sorted_types[:-1])}, and {sorted_types[-1]} samples"

def generate_shipment_document(shipment, form_data, document_type='invoice_packing'):
    """Generate document for a shipment
    
    Args:
        shipment: Shipment object
        form_data: Form data dictionary
        document_type: 'invoice_packing' for invoice & packing list only, 
                      'custom_docs' for full customs clearance documents
    """
    try:
        # Get the absolute path to our template file based on document type
        current_dir = os.path.dirname(os.path.abspath(__file__))
        root_dir = os.path.dirname(current_dir)
        
        # Determine template based on shipment type and document type
        if shipment.shipment_type == 'import':
            if document_type == 'invoice_packing':
                # For import shipments invoice & packing list, use Room Temperature invoice template
                template_path = os.path.join(root_dir, 'templates', 'Room Temperature Sample Import', 'invoice_packinglist_room_temperature.docx')
                doc_type_name = "room_temperature_import_invoice"
            else:  # custom_docs
                # For import shipments custom docs, use Room Temperature custom docs template
                template_path = os.path.join(root_dir, 'templates', 'Room Temperature Sample Import', 'Room_Temperature_Import_Custom_Docs .docx')
                doc_type_name = "room_temperature_import_custom"
        elif document_type == 'custom_docs':
            template_path = os.path.join(root_dir, 'templates', 'export_custom_docs.docx')
            doc_type_name = "custom_docs"
        else:  # default to invoice_packing for export/other shipment types
            template_path = os.path.join(root_dir, 'templates', 'invoice_packinglist.docx')
            doc_type_name = "invoice_packing"
        
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"Template file not found: {template_path}")
        
        # Load the template with docxtpl for initial replacements
        tpl = DocxTemplate(template_path)
        
        # Get aggregated sample types for this shipment
        sample_type_aggregated = get_aggregated_sample_types(form_data)
        
        # Get signing authority details
        def get_signing_authority_context():
            """Get signing authority context for document generation"""
            signing_authority = None
            if shipment.signing_authority_id:
                signing_authority = SigningAuthority.query.get(shipment.signing_authority_id)
            else:
                # Get default signing authority if none assigned
                signing_authority = SigningAuthority.query.filter_by(is_default=True, is_active=True).first()
            
            if signing_authority:
                return {
                    'signing_authority_name': signing_authority.name,
                    'signing_authority_designation': signing_authority.designation,
                    'singing_authority_department': signing_authority.department,  # Note: matches user's typo in placeholder
                    'signing_authority_organisation': signing_authority.organisation,
                    'signing_authority_contact_number': signing_authority.contact_number or '',
                    'signing_authority_contact_fax': signing_authority.contact_fax or '',
                    'signing_authority_email': signing_authority.email or '',
                    # Hindi placeholders
                    'signing_authority_name_hindi': signing_authority.name_hindi or '',
                    'signing_authority_designation_hindi': signing_authority.designation_hindi or '',
                    'signing_authority_department_hindi': signing_authority.department_hindi or ''
                }
            else:
                # Default empty values if no signing authority found
                return {
                    'signing_authority_name': '',
                    'signing_authority_designation': '',
                    'singing_authority_department': '',
                    'signing_authority_organisation': '',
                    'signing_authority_contact_number': '',
                    'signing_authority_contact_fax': '',
                    'signing_authority_email': '',
                    # Hindi placeholders
                    'signing_authority_name_hindi': '',
                    'signing_authority_designation_hindi': '',
                    'signing_authority_department_hindi': ''
                }
        
        # Get signing authority context
        signing_authority_context = get_signing_authority_context()
        
        # Prepare context based on shipment type
        if shipment.shipment_type == 'export':
            # Check if "other" consignee was selected
            consignee_selection = form_data.get('consignee', 'himadri')  # Default to himadri
            
            if consignee_selection == 'other':
                # Use other consignee details
                consignee_name = form_data.get('other_consignee_org', 'Other Organization')
                consignee_address = form_data.get('other_consignee_address', 'Other Address')
                consignee_contact = form_data.get('other_consignee_contact', 'Other Contact')
                consignee_phone = form_data.get('other_consignee_phone', 'Other Phone')
            else:
                # Use default Himadri details
                consignee_name = "Himadri - Indian Arctic Research Station C/O Kingsbay AS"
                consignee_address = "N-9173 Ny-Alesund, Norway"
                consignee_contact = "Kingsbay AS, Longyearbyen"
                consignee_phone = "+47 79 027200"
            
            # Set exporter details based on shipment type
            if shipment.shipment_type == 'import':
                # For import shipments, Himadri is the exporter (sending samples from Norway to India)
                exporter_name = "Himadri - Indian Arctic Research Station"
                exporter_ministry = ""  # Not applicable for Himadri
                exporter_address = "Ny-Alesund, c/o Kings Bay AS"
                exporter_country = "Norway"
                exporter_location = "N-9173, Ny-Alesund, Norway"
                exporter_phone = "+47 79 02 72 00"
                exporter_gst = ""  # Not applicable for Norwegian station
                exporter_lut = ""  # Not applicable for Norwegian station
            else:
                # For export/other shipments, NCPOR is the exporter (sending samples from India to Norway)
                exporter_name = "National Center for Polar and Ocean Research (NCPOR) [erstwhite NCAOR]"
                exporter_ministry = "Ministry of Earth Sciences (Govt. of India)"
                exporter_address = "Headland Sada, Vasco da Gama, Goa - 403804"
                exporter_country = "India"
                exporter_location = "India"
                exporter_phone = "+91 832 2525501"
                exporter_gst = "30AACFN4991P1ZN"
                exporter_lut = "AD300618000016R"
            
            context = {
                'invoice_no': shipment.invoice_number,
                'invoice_date': datetime.now().strftime('%d-%m-%Y'),
                'file_reference_number': shipment.file_reference_number or '',
                'Sample_Type': sample_type_aggregated,
                'ncpor_gst': "30AACFN4991P1ZN",
                'ncpor_lut': "AD300618000016R",
                'exporter_name': exporter_name,
                'exporter_ministry': exporter_ministry,
                'exporter_address': exporter_address,
                'exporter_country': exporter_country,
                'exporter_location': exporter_location,
                'exporter_phone': exporter_phone,
                'exporter_gst': exporter_gst,
                'exporter_lut': exporter_lut,
                'consignee_name': consignee_name,
                'consignee_address': consignee_address,
                'consignee_contact': consignee_contact,
                'consignee_phone': consignee_phone,
                'destination_country': form_data.get('destination_country', 'NORWAY'),
                'country_of_origin_goods': form_data.get('country_of_origin', 'India'),
                'country_of_final_destination': form_data.get('country_of_final_destination', 'Norway'),
                'airport_of_loading': form_data.get('airport_loading', 'MUMBAI'),
                'mode_of_transport': form_data.get('mode_of_transport', 'Air'),
                'transport_facility_type': 'Air Port' if form_data.get('mode_of_transport', 'Air') == 'Air' else 'Port',
                'port_of_loading': form_data.get('port_of_loading', ''),
                'port_of_discharge': form_data.get('port_of_discharge', ''),
                'requester_name': form_data.get('requester_name', ''),
                'expedition_year': form_data.get('expedition_year', ''),
                'batch_number': form_data.get('batch_number', ''),
                'return_type': form_data.get('return_type', ''),
                'purpose': "RESEARCH AND DEVELOPMENT",
                'declaration': "We declare that the invoice shows the actual price of goods described and that all particulars are true and correct.",
                **signing_authority_context  # Add signing authority details
            }
        elif shipment.shipment_type in ['import', 'reimport']:
            # Check if "other" consignee was selected
            consignee_selection = form_data.get('consignee', 'ncpor')  # Default to ncpor for import/reimport
            
            # Standard consignee placeholders ALWAYS contain NCPOR details for imports
            consignee_name = "National Center for Polar and Ocean Research (NCPOR) [erstwhite NCAOR]"
            consignee_address = "Headland Sada, Vasco da Gama, Goa - 403804"
            consignee_phone = "9274584406"
            consignee_gst = "30AACFN4991P1ZN"
            consignee_lut = "AD300618000016R"
            consignee_ministry = "Ministry of Earth Sciences (Govt. of India)"
            consignee_location = "India"
            
            context = {
                'invoice_no': shipment.invoice_number,
                'invoice_date': datetime.now().strftime('%d-%m-%Y'),
                'file_reference_number': shipment.file_reference_number or '',
                'Sample_Type': sample_type_aggregated,
                'ncpor_gst': "30AACFN4991P1ZN",
                'ncpor_lut': "AD300618000016R",
                'exporter_name': "Himadri - Indian Arctic Research Station",
                'exporter_org': "HIMADRI-Indian Arctic Research Station",
                'exporter_address': "Ny-Alesund, c/o Kings Bay AS",
                'exporter_location': "N-9173, Ny-Alesund, Norway",
                'exporter_phone': "+47 79 02 72 00",
                'consignee_name': consignee_name,
                'consignee_ministry': consignee_ministry,
                'consignee_address': consignee_address,
                'consignee_location': consignee_location,
                'consignee_phone': consignee_phone,
                'consignee_gst': consignee_gst,
                'consignee_lut': consignee_lut,
                'consignee_contact': "Dr. Rohit Srivastava",
                'consigner_name': "Himadri - Indian Arctic Research Station C/O Kingsbay AS",
                'consigner_address': "N-9173 Ny-Alesund, Norway",
                'consigner_contact': "Kingsbay AS, Longyearbyen",
                'consigner_phone': "+47 79 027200",
                'destination_country': 'INDIA',
                'country_of_origin_goods': form_data.get('country_of_origin', 'Norway'),
                'country_of_final_destination': form_data.get('country_of_final_destination', 'India'),
                'final_destination': form_data.get('final_destination', ''),
                'airport_of_loading': form_data.get('import_mode', 'AIR').upper(),
                'mode_of_transport': form_data.get('mode_of_transport', 'Air'),
                'transport_facility_type': 'Air Port' if form_data.get('mode_of_transport', 'Air') == 'Air' else 'Port',
                'port_of_loading': form_data.get('port_of_loading', ''),
                'port_of_discharge': form_data.get('port_of_discharge', ''),
                'requester_name': form_data.get('requester_name', ''),
                'expedition_year': form_data.get('expedition_year', ''),
                'batch_number': form_data.get('batch_number', ''),
                'import_purpose': form_data.get('import_purpose', 'RESEARCH'),
                'purpose': "RESEARCH AND DEVELOPMENT",
                'declaration': "We declare that the invoice shows the actual price of goods described and that all particulars are true and correct.",
                
                # Other Organization Details (populated when "other" consignee is selected)
                'other_org_name': form_data.get('other_consignee_org', '') if consignee_selection == 'other' else "",
                'other_org_ministry': form_data.get('other_consignee_ministry', '') if consignee_selection == 'other' else "",
                'other_org_address': form_data.get('other_consignee_address', '') if consignee_selection == 'other' else "",
                'other_org_location': form_data.get('other_consignee_country', '') if consignee_selection == 'other' else "",
                'other_org_phone': form_data.get('other_consignee_phone', '') if consignee_selection == 'other' else "",
                'other_org_gst': form_data.get('other_consignee_gst', '') if consignee_selection == 'other' else "",
                'other_org_lut': form_data.get('other_consignee_lut', '') if consignee_selection == 'other' else "",
                'other_org_contact': form_data.get('other_consignee_contact', '') if consignee_selection == 'other' else "",
                'other_org_designation': form_data.get('other_consignee_designation', '') if consignee_selection == 'other' else "",
                'other_org_email': form_data.get('other_consignee_email', '') if consignee_selection == 'other' else "",
                
                **signing_authority_context  # Add signing authority details
            }
        else:  # cold
            context = {
                'invoice_no': shipment.invoice_number,
                'invoice_date': datetime.now().strftime('%d-%m-%Y'),
                'file_reference_number': shipment.file_reference_number or '',
                'Sample_Type': sample_type_aggregated,
                'requester_name': form_data.get('requester_name', ''),
                'expedition_year': form_data.get('expedition_year', ''),
                'batch_number': form_data.get('batch_number', ''),
                'total_boxes': form_data.get('total_boxes', '0'),
                'purpose': "RESEARCH AND DEVELOPMENT",
                'declaration': "We declare that the invoice shows the actual price of goods described and that all particulars are true and correct.",
                **signing_authority_context  # Add signing authority details
            }
        
        # Render the template with docxtpl
        tpl.render(context)
        
        # Save the docxtpl output to a temporary file
        if shipment.shipment_type == 'import':
            temp_path = os.path.join(root_dir, 'templates', 'temp_room_temperature_import.docx')
        else:
            temp_path = os.path.join(root_dir, 'templates', f'temp_{shipment.shipment_type}_{doc_type_name}.docx')
        tpl.save(temp_path)
        
        # Load the temporary file with python-docx for table manipulation
        doc = Document(temp_path)
        
        # Handle table placement and population
        invoice_table = handle_table_placement(doc, form_data)
        total_amount = populate_table_data(invoice_table, form_data)

        # Handle packing list table
        pl_table = handle_pl_table_placement(doc, form_data)
        populate_pl_table_data(pl_table, form_data)
        
        # Handle shipper table
        shipper_table = handle_shipper_table_placement(doc, form_data)
        populate_shipper_table_data(shipper_table, form_data)
        
        # Compute amount in words
        amount_in_words = f"USD {num2words(int(total_amount))} Only"

        # Update amount in words and total amount
        for paragraph in doc.paragraphs:
            if '[AMOUNT_IN_WORDS]' in paragraph.text:
                paragraph.text = paragraph.text.replace('[AMOUNT_IN_WORDS]', amount_in_words)
            if '[TOTAL_AMOUNT]' in paragraph.text:
                paragraph.text = paragraph.text.replace('[TOTAL_AMOUNT]', f"{total_amount:.0f}")
        
        # Save to buffer
        docx_buffer = io.BytesIO()
        doc.save(docx_buffer)
        docx_buffer.seek(0)
        
        # Remove the temporary file
        os.remove(temp_path)
        
        # Generate filename using new format
        from .utils.helpers import generate_document_filename
        filename = generate_document_filename(shipment, form_data, document_type)
        
        # Return the generated document
        return send_file(
            docx_buffer,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=filename
        )
        
    except Exception as e:
        raise e

@main.route('/dashboard')
@login_required
def dashboard():
    if current_user.is_admin():
        # Admin dashboard - show all shipments from PI and Field Personnel
        all_shipments = Shipment.query.join(User, Shipment.created_by == User.id).order_by(Shipment.created_at.desc()).all()
        users_count = User.query.count()
        admin_users = User.query.join(User.roles).filter(Role.name == 'Admin').all()
        non_admin_users = User.query.filter(~User.id.in_([u.id for u in admin_users])).all()
        signing_authorities = SigningAuthority.query.filter_by(is_active=True).order_by(SigningAuthority.is_default.desc(), SigningAuthority.name.asc()).all()
        
        return render_template('dashboard/admin_dashboard.html', 
                             user=current_user, 
                             shipments=all_shipments,
                             users_count=users_count,
                             admin_users=admin_users,
                             non_admin_users=non_admin_users,
                             signing_authorities=signing_authorities)
    else:
        # Regular user dashboard - show only their shipments
        user_shipments = Shipment.query.filter_by(created_by=current_user.id).order_by(Shipment.created_at.desc()).all()
        return render_template('dashboard/user_dashboard.html', 
                             user=current_user, 
                             shipments=user_shipments)

@main.route('/admin/users')
@login_required
@admin_required
def admin_users():
    """Admin page to manage users"""
    users = User.query.order_by(User.created_at.desc()).all()
    roles = Role.query.all()
    return render_template('admin/users.html', users=users, roles=roles)

@main.route('/admin/create-user', methods=['POST'])
@login_required
@admin_required
def admin_create_user():
    """Admin endpoint to create new users"""
    email = request.form.get('email')
    password = request.form.get('password')
    first_name = request.form.get('first_name')
    last_name = request.form.get('last_name')
    phone = request.form.get('phone')
    organization = request.form.get('organization')
    role_ids = request.form.getlist('roles')

    if not all([email, password, first_name, last_name, phone, organization]):
        flash('All fields are required', 'error')
        return redirect(url_for('main.admin_users'))

    if User.query.filter_by(email=email).first():
        flash('Email already registered', 'error')
        return redirect(url_for('main.admin_users'))

    # Create new user
    new_user = User(
        email=email,
        password=generate_password_hash(password),
        first_name=first_name,
        last_name=last_name,
        phone=phone,
        organization=organization,
        unique_id=User.generate_unique_id()
    )

    # Assign roles (prevent duplicates and enforce role separation)
    selected_roles = []
    is_admin_selected = False
    
    for role_id in role_ids:
        role = Role.query.get(role_id)
        if role and role not in new_user.roles:
            selected_roles.append(role)
            if role.name == 'Admin':
                is_admin_selected = True
    
    # Enforce role separation: Admins cannot have PI or Field Personnel roles
    if is_admin_selected:
        # If Admin role is selected, only assign Admin role
        admin_role = next((role for role in selected_roles if role.name == 'Admin'), None)
        if admin_role:
            new_user.roles.append(admin_role)
    else:
        # If Admin role is not selected, assign other roles
        for role in selected_roles:
            if role.name != 'Admin':  # Extra safety check
                new_user.roles.append(role)

    try:
        db.session.add(new_user)
        db.session.commit()
        flash(f'User {email} created successfully!', 'success')
    except Exception as e:
        db.session.rollback()
        if 'UNIQUE constraint failed: user_roles' in str(e):
            flash('Error: Duplicate role assignment detected. Please try again or check if user already exists.', 'error')
        else:
            flash(f'Error creating user: {str(e)}', 'error')
    
    return redirect(url_for('main.admin_users'))

@main.route('/admin/toggle-user/<int:user_id>')
@login_required
@admin_required
def admin_toggle_user(user_id):
    """Admin endpoint to activate/deactivate users"""
    user = User.query.get_or_404(user_id)
    
    if user.id == current_user.id:
        flash('You cannot deactivate your own account', 'error')
        return redirect(url_for('main.admin_users'))
    
    user.is_active = not user.is_active
    db.session.commit()
    
    status = 'activated' if user.is_active else 'deactivated'
    flash(f'User {user.email} has been {status}', 'success')
    return redirect(url_for('main.admin_users'))

@main.route('/admin/edit-user/<int:user_id>')
@login_required
@admin_required
def admin_edit_user(user_id):
    """Admin endpoint to edit user details"""
    user = User.query.get_or_404(user_id)
    roles = Role.query.all()
    return render_template('admin/edit_user.html', user=user, roles=roles)

@main.route('/admin/update-user/<int:user_id>', methods=['POST'])
@login_required
@admin_required
def admin_update_user(user_id):
    """Admin endpoint to update user details"""
    user = User.query.get_or_404(user_id)
    
    # Get form data
    email = request.form.get('email')
    first_name = request.form.get('first_name')
    last_name = request.form.get('last_name')
    phone = request.form.get('phone')
    organization = request.form.get('organization')
    new_password = request.form.get('new_password')
    role_ids = request.form.getlist('roles')

    # Validate required fields
    if not all([email, first_name, last_name, phone, organization]):
        flash('All fields except password are required', 'error')
        return redirect(url_for('main.admin_edit_user', user_id=user_id))

    # Check if email is taken by another user
    if email != user.email:
        existing_user = User.query.filter_by(email=email).first()
        if existing_user:
            flash('Email address is already taken by another user', 'error')
            return redirect(url_for('main.admin_edit_user', user_id=user_id))

    # Update user basic info
    user.email = email
    user.first_name = first_name
    user.last_name = last_name
    user.phone = phone
    user.organization = organization

    # Update password if provided
    if new_password and len(new_password.strip()) > 0:
        user.password = generate_password_hash(new_password)

    # Update roles with same logic as create user
    selected_roles = []
    is_admin_selected = False
    
    for role_id in role_ids:
        role = Role.query.get(role_id)
        if role:
            selected_roles.append(role)
            if role.name == 'Admin':
                is_admin_selected = True
    
    # Clear existing roles
    user.roles.clear()
    
    # Apply role separation logic
    if is_admin_selected:
        # If Admin role is selected, only assign Admin role
        admin_role = next((role for role in selected_roles if role.name == 'Admin'), None)
        if admin_role:
            user.roles.append(admin_role)
    else:
        # If Admin role is not selected, assign other roles
        for role in selected_roles:
            if role.name != 'Admin':  # Extra safety check
                user.roles.append(role)

    try:
        db.session.commit()
        flash(f'User {user.email} updated successfully!', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'Error updating user: {str(e)}', 'error')
    
    return redirect(url_for('main.admin_users'))

@main.route('/admin/delete-user/<int:user_id>')
@login_required
@admin_required
def admin_delete_user(user_id):
    """Admin endpoint to delete a user"""
    user = User.query.get_or_404(user_id)
    
    # Prevent deleting own account
    if user.id == current_user.id:
        flash('You cannot delete your own account', 'error')
        return redirect(url_for('main.admin_users'))
    
    # Check if user has any shipments
    user_shipments = Shipment.query.filter_by(created_by=user.id).count()
    
    if user_shipments > 0:
        flash(f'Cannot delete user {user.email}: User has {user_shipments} shipment(s). Please delete or reassign shipments first.', 'error')
        return redirect(url_for('main.admin_users'))
    
    # Store user details for confirmation message
    user_email = user.email
    user_name = f"{user.first_name} {user.last_name}"
    
    try:
        # Delete the user completely from database
        db.session.delete(user)
        db.session.commit()
        
        flash(f'User {user_name} ({user_email}) has been permanently deleted!', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'Error deleting user: {str(e)}', 'error')
    
    return redirect(url_for('main.admin_users'))

@main.route('/admin/qr-codes')
@login_required
@admin_required
def qr_codes_management():
    """Admin page to manage QR codes for all packages"""
    # Get all packages with QR codes, with pagination
    page = request.args.get('page', 1, type=int)
    per_page = 20
    
    # Get shipment filter
    shipment_filter = request.args.get('shipment_id', type=int)
    status_filter = request.args.get('status', '')
    
    query = PackageQRCode.query.join(Shipment)
    
    if shipment_filter:
        query = query.filter(PackageQRCode.shipment_id == shipment_filter)
    
    if status_filter:
        query = query.filter(Shipment.status == status_filter)
    
    packages = query.order_by(PackageQRCode.created_at.desc()).paginate(
        page=page, per_page=per_page, error_out=False
    )
    
    # Get all shipments for filter dropdown
    all_shipments = Shipment.query.order_by(Shipment.created_at.desc()).limit(100).all()
    
    # Get status options
    status_options = ['Acknowledged', 'Document_Generation', 'Quotation_Requested', 'Awaiting_Quotation_Approval']
    
    return render_template('admin/qr_codes.html', 
                         packages=packages, 
                         all_shipments=all_shipments,
                         status_options=status_options,
                         current_shipment_filter=shipment_filter,
                         current_status_filter=status_filter)

@main.route('/admin/regenerate-qr/<int:package_id>')
@login_required
@admin_required
def regenerate_qr_code(package_id):
    """Regenerate QR code for a specific package"""
    package_qr = PackageQRCode.query.get_or_404(package_id)
    
    try:
        qr_service = QRCodeService()
        base_url = request.url_root.rstrip('/')
        
        updated_package = qr_service.regenerate_qr_code(package_qr, base_url)
        
        if updated_package:
            flash(f'QR code regenerated successfully for package #{package_qr.package_number} in shipment {package_qr.shipment.invoice_number}', 'success')
        else:
            flash('Failed to regenerate QR code. Please try again.', 'error')
            
    except Exception as e:
        current_app.logger.error(f"Error regenerating QR code for package {package_id}: {str(e)}")
        flash('Error occurred while regenerating QR code.', 'error')
    
    return redirect(url_for('main.qr_codes_management'))

@main.route('/admin/download-qr/<int:package_id>')
@login_required
@admin_required
def download_qr_code(package_id):
    """Download QR code image for a specific package"""
    package_qr = PackageQRCode.query.get_or_404(package_id)
    
    if not package_qr.qr_image_path:
        flash('QR code image not found. Try regenerating the QR code.', 'error')
        return redirect(url_for('main.qr_codes_management'))
    
    try:
        qr_file_path = os.path.join(current_app.root_path, package_qr.qr_image_path)
        
        if os.path.exists(qr_file_path):
            filename = f"qr_code_{package_qr.shipment.invoice_number.replace('/', '_')}_pkg_{package_qr.package_number}.png"
            return send_file(qr_file_path, as_attachment=True, download_name=filename)
        else:
            flash('QR code file not found. Try regenerating the QR code.', 'error')
            
    except Exception as e:
        current_app.logger.error(f"Error downloading QR code for package {package_id}: {str(e)}")
        flash('Error occurred while downloading QR code.', 'error')
    
    return redirect(url_for('main.qr_codes_management'))

@main.route('/admin/qr-bulk-actions', methods=['POST'])
@login_required
@admin_required
def qr_bulk_actions():
    """Handle bulk actions for QR codes"""
    action = request.form.get('action')
    package_ids = request.form.getlist('package_ids')
    
    if not package_ids:
        flash('No packages selected.', 'error')
        return redirect(url_for('main.qr_codes_management'))
    
    package_ids = [int(pid) for pid in package_ids if pid.isdigit()]
    
    if action == 'regenerate':
        try:
            qr_service = QRCodeService()
            base_url = request.url_root.rstrip('/')
            success_count = 0
            
            for package_id in package_ids:
                package_qr = PackageQRCode.query.get(package_id)
                if package_qr:
                    updated = qr_service.regenerate_qr_code(package_qr, base_url)
                    if updated:
                        success_count += 1
            
            flash(f'Successfully regenerated {success_count} out of {len(package_ids)} QR codes.', 'success')
            
        except Exception as e:
            current_app.logger.error(f"Error in bulk QR regeneration: {str(e)}")
            flash('Error occurred during bulk regeneration.', 'error')
    
    elif action == 'cleanup':
        try:
            qr_service = QRCodeService()
            removed_count = qr_service.cleanup_orphaned_qr_codes()
            flash(f'Cleaned up {removed_count} orphaned QR code files.', 'success')
            
        except Exception as e:
            current_app.logger.error(f"Error in QR cleanup: {str(e)}")
            flash('Error occurred during cleanup.', 'error')
    
    return redirect(url_for('main.qr_codes_management'))

@main.route('/admin/webapp-qr')
@login_required
@admin_required
def webapp_qr_code():
    """Generate and display QR code for the web app homepage"""
    try:
        qr_service = QRCodeService()
        base_url = request.url_root.rstrip('/')
        
        # Generate QR code for homepage
        qr_image_path = qr_service.generate_webapp_qr_code(base_url)
        
        # Get QR code URL for display
        qr_code_url = f"{base_url}/static/qrcodes/webapp_qr.png"
        
        return render_template('admin/webapp_qr.html', 
                             qr_code_url=qr_code_url,
                             webapp_url=base_url)
                             
    except Exception as e:
        import traceback
        error_details = traceback.format_exc()
        current_app.logger.error(f"Error generating webapp QR code: {str(e)}\n{error_details}")
        flash(f'Error generating web app QR code: {str(e)}', 'error')
        return redirect(url_for('main.dashboard'))

@main.route('/admin/download-webapp-qr')
@login_required
@admin_required
def download_webapp_qr():
    """Download the web app QR code image"""
    try:
        qr_codes_dir = os.path.join(current_app.root_path, 'static', 'qrcodes')
        qr_image_path = os.path.join(qr_codes_dir, 'webapp_qr.png')
        
        if os.path.exists(qr_image_path):
            return send_file(qr_image_path, 
                           as_attachment=True, 
                           download_name='COMPASS_WebApp_QR.png',
                           mimetype='image/png')
        else:
            flash('Web app QR code not found. Please generate it first.', 'error')
            return redirect(url_for('main.webapp_qr_code'))
            
    except Exception as e:
        current_app.logger.error(f"Error downloading webapp QR code: {str(e)}")
        flash('Error downloading QR code.', 'error')
        return redirect(url_for('main.webapp_qr_code'))

@main.route('/admin/signing-authorities')
@login_required
@admin_required
def admin_signing_authorities():
    """Admin page to manage signing authorities"""
    signing_authorities = SigningAuthority.query.order_by(SigningAuthority.created_at.desc()).all()
    return render_template('admin/signing_authorities.html', signing_authorities=signing_authorities)

@main.route('/admin/create-signing-authority', methods=['POST'])
@login_required
@admin_required
def admin_create_signing_authority():
    """Admin endpoint to create new signing authority"""
    name = request.form.get('name')
    designation = request.form.get('designation')
    department = request.form.get('department')
    organisation = request.form.get('organisation')
    
    # Hindi fields
    name_hindi = request.form.get('name_hindi')
    designation_hindi = request.form.get('designation_hindi')
    department_hindi = request.form.get('department_hindi')
    
    contact_number = request.form.get('contact_number')
    contact_fax = request.form.get('contact_fax')
    email = request.form.get('email')
    is_default = request.form.get('is_default') == 'on'

    if not all([name, designation, department, organisation]):
        flash('Name, designation, department, and organisation are required', 'error')
        return redirect(url_for('main.admin_signing_authorities'))

    # If this is being set as default, unset other defaults
    if is_default:
        SigningAuthority.query.update({'is_default': False})

    # Create new signing authority
    new_authority = SigningAuthority(
        name=name,
        designation=designation,
        department=department,
        organisation=organisation,
        name_hindi=name_hindi,
        designation_hindi=designation_hindi,
        department_hindi=department_hindi,
        contact_number=contact_number,
        contact_fax=contact_fax,
        email=email,
        is_default=is_default,
        created_by=current_user.id
    )

    try:
        db.session.add(new_authority)
        db.session.commit()
        flash(f'Signing Authority "{name}" created successfully!', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'Error creating signing authority: {str(e)}', 'error')
    
    return redirect(url_for('main.admin_signing_authorities'))

@main.route('/admin/edit-signing-authority/<int:authority_id>')
@login_required
@admin_required
def admin_edit_signing_authority(authority_id):
    """Admin endpoint to edit signing authority details"""
    authority = SigningAuthority.query.get_or_404(authority_id)
    return render_template('admin/edit_signing_authority.html', authority=authority)

@main.route('/admin/update-signing-authority/<int:authority_id>', methods=['POST'])
@login_required
@admin_required
def admin_update_signing_authority(authority_id):
    """Admin endpoint to update signing authority details"""
    authority = SigningAuthority.query.get_or_404(authority_id)
    
    # Get form data
    name = request.form.get('name')
    designation = request.form.get('designation')
    department = request.form.get('department')
    organisation = request.form.get('organisation')
    
    # Hindi fields
    name_hindi = request.form.get('name_hindi')
    designation_hindi = request.form.get('designation_hindi')
    department_hindi = request.form.get('department_hindi')
    
    contact_number = request.form.get('contact_number')
    contact_fax = request.form.get('contact_fax')
    email = request.form.get('email')
    is_default = request.form.get('is_default') == 'on'

    # Validate required fields
    if not all([name, designation, department, organisation]):
        flash('Name, designation, department, and organisation are required', 'error')
        return redirect(url_for('main.admin_edit_signing_authority', authority_id=authority_id))

    # If this is being set as default, unset other defaults
    if is_default:
        SigningAuthority.query.filter(SigningAuthority.id != authority_id).update({'is_default': False})

    # Update authority details
    authority.name = name
    authority.designation = designation
    authority.department = department
    authority.organisation = organisation
    authority.name_hindi = name_hindi
    authority.designation_hindi = designation_hindi
    authority.department_hindi = department_hindi
    authority.contact_number = contact_number
    authority.contact_fax = contact_fax
    authority.email = email
    authority.is_default = is_default

    try:
        db.session.commit()
        flash(f'Signing Authority "{authority.name}" updated successfully!', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'Error updating signing authority: {str(e)}', 'error')
    
    return redirect(url_for('main.admin_signing_authorities'))

@main.route('/admin/toggle-signing-authority/<int:authority_id>')
@login_required
@admin_required
def admin_toggle_signing_authority(authority_id):
    """Admin endpoint to activate/deactivate signing authorities"""
    authority = SigningAuthority.query.get_or_404(authority_id)
    
    authority.is_active = not authority.is_active
    
    # If deactivating the default authority, unset default flag
    if not authority.is_active and authority.is_default:
        authority.is_default = False
    
    db.session.commit()
    
    status = 'activated' if authority.is_active else 'deactivated'
    flash(f'Signing Authority "{authority.name}" has been {status}', 'success')
    return redirect(url_for('main.admin_signing_authorities'))

@main.route('/admin/delete-signing-authority/<int:authority_id>')
@login_required
@admin_required
def admin_delete_signing_authority(authority_id):
    """Admin endpoint to delete a signing authority"""
    authority = SigningAuthority.query.get_or_404(authority_id)
    
    # Check if authority is used in any shipments
    shipments_count = Shipment.query.filter_by(signing_authority_id=authority.id).count()
    
    if shipments_count > 0:
        flash(f'Cannot delete signing authority "{authority.name}": It is being used by {shipments_count} shipment(s).', 'error')
        return redirect(url_for('main.admin_signing_authorities'))
    
    # Store authority details for confirmation message
    authority_name = authority.name
    
    try:
        # Delete the authority completely from database
        db.session.delete(authority)
        db.session.commit()
        
        flash(f'Signing Authority "{authority_name}" has been permanently deleted!', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'Error deleting signing authority: {str(e)}', 'error')
    
    return redirect(url_for('main.admin_signing_authorities'))

@main.route('/admin/assign-signing-authority', methods=['POST'])
@login_required
@admin_required
def admin_assign_signing_authority():
    """Admin endpoint to assign signing authority to a shipment"""
    try:
        shipment_id = request.form.get('shipment_id')
        authority_id = request.form.get('authority_id')
        
        if not shipment_id or not authority_id:
            return jsonify({'success': False, 'message': 'Missing shipment ID or authority ID'})
        
        shipment = Shipment.query.get_or_404(shipment_id)
        authority = SigningAuthority.query.get_or_404(authority_id)
        
        # Update shipment with signing authority
        shipment.signing_authority_id = authority_id
        db.session.commit()
        
        return jsonify({
            'success': True, 
            'message': f'Signing Authority "{authority.name}" assigned successfully',
            'authority_name': authority.name,
            'authority_designation': authority.designation
        })
        
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)})

@main.route('/admin/acknowledge-shipment/<int:shipment_id>')
@login_required
@admin_required
def admin_acknowledge_shipment(shipment_id):
    """Admin endpoint to acknowledge receipt of a shipment"""
    from .utils.helpers import generate_file_reference_number
    
    shipment = Shipment.query.get_or_404(shipment_id)
    
    # Allow acknowledgment regardless of current status
    # Update shipment status to acknowledged (or keep current status if it's more advanced)
    if shipment.status not in ['Acknowledged', 'Document_Generated', 'Delivered']:
        shipment.status = 'Acknowledged'
    
    # Set acknowledgment details
    shipment.acknowledged_by = current_user.id
    shipment.acknowledged_at = datetime.now()
    
    # Generate file reference number if not already generated
    if not shipment.file_reference_number:
        shipment.file_reference_number = generate_file_reference_number(shipment, current_user)
    
    db.session.commit()
    
    flash(f'Shipment {shipment.invoice_number} acknowledged successfully! File Reference: {shipment.file_reference_number}', 'success')
    return redirect(url_for('main.dashboard'))

@main.route('/admin/generate-document/<int:shipment_id>')
@main.route('/admin/generate-document/<int:shipment_id>/<document_type>')
@login_required
@admin_required
def admin_generate_document(shipment_id, document_type='invoice_packing'):
    """Admin endpoint to generate document for a shipment - available regardless of status"""
    shipment = Shipment.query.get_or_404(shipment_id)
    
    # Validate document type
    if document_type not in ['invoice_packing', 'custom_docs']:
        document_type = 'invoice_packing'  # default fallback
    
    # Allow document generation regardless of current status
    try:
        # Parse the stored form data
        form_data = json.loads(shipment.form_data)
        
        # Update status to document generated (or keep if already delivered)
        if shipment.status != 'Delivered':
            shipment.status = 'Document_Generated'
        
        # Auto-acknowledge if not already acknowledged
        if not shipment.acknowledged_by:
            shipment.acknowledged_by = current_user.id
            shipment.acknowledged_at = datetime.now()
            # Generate file reference number if not already generated
            if not shipment.file_reference_number:
                shipment.file_reference_number = generate_file_reference_number(shipment, current_user)
        
        db.session.commit()
        
        # Generate and return the document with specified type
        return generate_shipment_document(shipment, form_data, document_type)
        
    except Exception as e:
        shipment.status = 'Failed'
        db.session.commit()
        flash(f'Error generating document: {str(e)}', 'error')
        return redirect(url_for('main.dashboard'))

@main.route('/admin/add-comment', methods=['POST'])
@login_required
@admin_required
def admin_add_comment():
    """Admin endpoint to add comments to shipments"""
    try:
        shipment_id = request.form.get('shipment_id')
        comment = request.form.get('comment')
        
        if not shipment_id or not comment:
            return jsonify({'success': False, 'message': 'Missing shipment ID or comment'})
        
        shipment = Shipment.query.get_or_404(shipment_id)
        
        # Update shipment with comment
        shipment.admin_comment = comment
        shipment.comment_by = current_user.id
        shipment.comment_at = datetime.now()
        
        # If comment indicates changes needed, update status
        if any(keyword in comment.lower() for keyword in ['change', 'modify', 'update', 'correct', 'fix']):
            shipment.status = 'Needs_Changes'
        
        db.session.commit()
        
        return jsonify({'success': True, 'message': 'Comment added successfully'})
        
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)})

@main.route('/admin/combine-shipments', methods=['POST'])
@login_required
@admin_required
def admin_combine_shipments():
    """Admin endpoint to combine multiple shipments"""
    try:
        shipment_ids_json = request.form.get('shipment_ids')
        shipment_ids = json.loads(shipment_ids_json)
        
        if len(shipment_ids) < 2:
            flash('At least 2 shipments are required for combining', 'error')
            return redirect(url_for('main.dashboard'))
        
        # Get all shipments to combine
        shipments = Shipment.query.filter(Shipment.id.in_(shipment_ids)).all()
        
        if len(shipments) != len(shipment_ids):
            flash('Some shipments could not be found', 'error')
            return redirect(url_for('main.dashboard'))
        
        # Validate all shipments are of same type and year
        first_shipment = shipments[0]
        if not all(s.shipment_type == first_shipment.shipment_type and 
                  s.expedition_year == first_shipment.expedition_year for s in shipments):
            flash('Can only combine shipments of the same type and expedition year', 'error')
            return redirect(url_for('main.dashboard'))
        
        # Store shipment IDs in session for the combine form
        from flask import session
        session['combine_shipment_ids'] = shipment_ids
        
        # Redirect to combine form
        return redirect(url_for('main.admin_combine_form'))
        
    except Exception as e:
        flash(f'Error processing shipments for combination: {str(e)}', 'error')
        return redirect(url_for('main.dashboard'))

@main.route('/admin/combine-form', methods=['GET', 'POST'])
@login_required
@admin_required
def admin_combine_form():
    """Admin page to review and edit combined shipment before generation"""
    from flask import session
    
    # Handle POST request from dashboard (new direct route)
    if request.method == 'POST':
        shipment_ids = request.form.getlist('shipment_ids')
        print(f"DEBUG: Received shipment_ids: {shipment_ids}")
        
        if not shipment_ids:
            flash('No shipments selected for combining', 'error')
            return redirect(url_for('main.dashboard'))
        
        # Convert to integers and store in session
        try:
            shipment_ids = [int(sid) for sid in shipment_ids if sid]  # Filter out empty strings
            print(f"DEBUG: Converted shipment_ids: {shipment_ids}")
            session['combine_shipment_ids'] = shipment_ids
        except (ValueError, TypeError) as e:
            print(f"DEBUG: Error converting shipment_ids: {e}")
            flash('Invalid shipment IDs provided', 'error')
            return redirect(url_for('main.dashboard'))
    
    # Get shipment IDs from session (for both GET and POST after storing)
    shipment_ids = session.get('combine_shipment_ids', [])
    print(f"DEBUG: shipment_ids from session: {shipment_ids}")
    
    # DEBUG: If no shipments in session, auto-select test shipments for debugging
    if not shipment_ids:
        # Try to get some test export shipments for debugging
        test_shipments = Shipment.query.filter_by(
            shipment_type='export',
            status='Submitted'
        ).filter(
            Shipment.is_combined == False
        ).limit(3).all()
        
        if test_shipments:
            shipment_ids = [s.id for s in test_shipments]
            session['combine_shipment_ids'] = shipment_ids
            flash(f'DEBUG: Auto-selected {len(shipment_ids)} test shipments for combining', 'info')
            print(f"DEBUG: Auto-selected test shipments: {shipment_ids}")
        else:
            flash('No shipments available for combining. Create test data first!', 'error')
        return redirect(url_for('main.dashboard'))
    
    # Get all shipments to combine
    try:
        shipments = Shipment.query.filter(Shipment.id.in_(shipment_ids)).all()
        print(f"DEBUG: Found {len(shipments)} shipments")
        
        if len(shipments) != len(shipment_ids):
            flash('Some selected shipments could not be found', 'error')
            return redirect(url_for('main.dashboard'))
        
        # Validate shipments are compatible
        if len(shipments) < 2:
            flash('At least 2 shipments are required for combining', 'error')
            return redirect(url_for('main.dashboard'))
        
    except Exception as e:
        print(f"DEBUG: Error querying shipments: {e}")
        flash('Error retrieving shipments for combining', 'error')
        return redirect(url_for('main.dashboard'))
    
    # Get next unique combined number for display (not used in new format, just for fallback)
    try:
        combined_number = CombinedShipmentCounter.get_next_number()
        print(f"DEBUG: Combined number: {combined_number}")
    except Exception as e:
        print(f"DEBUG: Error getting combined number: {e}")
        # Use timestamp as fallback instead of incrementing counter
        import time
        combined_number = int(time.time()) % 10000  # Last 4 digits of timestamp
    
    # Generate combined invoice number using the NEW logic with serial numbers
    first_shipment = shipments[0]
    
    print(f"DEBUG: Starting invoice generation for {len(shipments)} shipments")
    print(f"DEBUG: First shipment type: {first_shipment.shipment_type}")
    
    # Get current date for proper format
    from datetime import datetime
    current_date = datetime.now()
    year = current_date.strftime('%Y')
    month = current_date.strftime('%b').upper()  # JAN, FEB, MAR, etc.
    
    print(f"DEBUG: Date components - Year: {year}, Month: {month}")
    
    # Get next serial number for the combined shipment
    try:
        serial_number = ShipmentSerialCounter.get_next_serial()
        print(f"DEBUG: Generated serial number: {serial_number}")
    except Exception as e:
        print(f"DEBUG: Error getting serial number: {e}")
        serial_number = "0001"  # Fallback
    
    # Collect unique IDs from all shipments being combined
    unique_user_ids = []
    seen_users = set()
    
    print(f"DEBUG: Collecting unique user IDs...")
    for i, shipment in enumerate(shipments):
        print(f"DEBUG: Processing shipment {i+1}: {shipment.invoice_number}")
        user_id = shipment.created_by
        print(f"DEBUG: Shipment {i+1} created_by: {user_id}")
        
        if user_id not in seen_users:
            try:
                user = User.query.get(user_id)
                if user:
                    print(f"DEBUG: Found user: {user.first_name} {user.last_name}, unique_id: {user.unique_id}")
                    if user.unique_id:
                        unique_user_ids.append(user.unique_id)
                        seen_users.add(user_id)
                    else:
                        print(f"DEBUG: WARNING - User {user.first_name} {user.last_name} has no unique_id!")
                else:
                    print(f"DEBUG: WARNING - User with ID {user_id} not found!")
            except Exception as e:
                print(f"DEBUG: Error getting user {user_id}: {e}")
    
    # Ensure we have at least one unique ID
    if not unique_user_ids:
        print(f"DEBUG: WARNING - No unique IDs found, using fallback")
        unique_user_ids = ["UNKNOWN"]
    
    # Create combined unique IDs string
    unique_ids_str = "/".join(unique_user_ids)
    print(f"DEBUG: Combined unique IDs: {unique_ids_str}")
    
    # Generate invoice number following the correct format: NCPOR/ARC/YYYY/MMM/EXP/TYPE/CMB/Unique_IDs/Serial
    # Since we're combining shipments, we always use CMB (combined) logic similar to real-time method
    if first_shipment.shipment_type == 'export':
        # Extract return type from ALL shipments being combined to determine the type
        return_types = set()
        print(f"DEBUG: Analyzing return types from all shipments...")
        
        for shipment in shipments:
            print(f"DEBUG: Analyzing invoice: {shipment.invoice_number}")
            original_invoice_parts = shipment.invoice_number.split('/')
            print(f"DEBUG: Invoice parts: {original_invoice_parts}")
            
            if len(original_invoice_parts) >= 6:
                return_type_part = original_invoice_parts[5]
                return_types.add(return_type_part)
                print(f"DEBUG: Found return type: {return_type_part}")
            else:
                print(f"DEBUG: Invoice format doesn't have return type at position 5")
        
        # If all shipments have the same return type, use it; otherwise use RET as default
        if len(return_types) == 1:
            return_type = return_types.pop()
            print(f"DEBUG: All shipments have same return type: {return_type}")
        else:
            return_type = 'RET'  # Default when mixed types
            print(f"DEBUG: Mixed or no return types found, using default: {return_type}")
        
        # Use CMB format like real-time method since we're combining multiple shipments
        combined_invoice = f"NCPOR/ARC/{year}/{month}/EXP/{return_type}/CMB/{unique_ids_str}/{serial_number}"
        
    elif first_shipment.shipment_type == 'import':
        # Use new NCPOR/ARC pattern for import combined shipments
        combined_invoice = f"NCPOR/ARC/{year}/{month}/SAM/RT/CMB/{unique_ids_str}/{serial_number}"
        
    elif first_shipment.shipment_type == 'reimport':
        # Use CMB format for reimport combined shipments too
        combined_invoice = f"NCPOR/REIMP/{year}/{month}/RESEARCH/CMB/{unique_ids_str}/{serial_number}"
        
    else:  # cold
        # Use CMB format for cold combined shipments too
        combined_invoice = f"NCPOR/COLD/{year}/{month}/CMB/{unique_ids_str}/{serial_number}"
        
    print(f"DEBUG: Successfully generated combined invoice: {combined_invoice}")
    
    # Combine form data from all shipments
    combined_packages = []
    total_packages = 0
    
    try:
        for i, shipment in enumerate(shipments):
            print(f"DEBUG: Processing shipment {i+1}: {shipment.invoice_number}")
            
            form_data = json.loads(shipment.form_data) if shipment.form_data else {}
            current_packages = int(form_data.get('total_packages', 0))
            print(f"DEBUG: Shipment {i+1} has {current_packages} packages")
            
            # Extract package and item data
            for pkg_num in range(1, current_packages + 1):
                new_pkg_num = total_packages + pkg_num
                package_data = {
                    'package_number': new_pkg_num,
                    'type': get_package_type_display_name(form_data.get(f'package_{pkg_num}_type', ''), form_data, pkg_num),
                    'length': form_data.get(f'package_{pkg_num}_length', ''),
                    'width': form_data.get(f'package_{pkg_num}_width', ''),
                    'height': form_data.get(f'package_{pkg_num}_height', ''),
                    'item_list': [],  # Changed from 'items' to avoid conflict with dict.items()
                    'items_count': 0,  # Pre-calculated count
                    'source_shipment': shipment.invoice_number,
                    'source_user': f"{shipment.created_by_user.first_name} {shipment.created_by_user.last_name}"
                }
                
                # Get items for this package
                items_count = int(form_data.get(f'package_{pkg_num}_items_count', 0))
                print(f"DEBUG: Package {pkg_num} has {items_count} items")
                
                # Create items list
                items_list = []
                for item_num in range(1, items_count + 1):
                    prefix = f'package_{pkg_num}_item_{item_num}'
                    # Use original shipment's requester name as default for attn field
                    original_requester = shipment.requester_name or 'Unknown'
                    item_data = {
                        'description': form_data.get(f'{prefix}_description', ''),
                        'hsn_code': form_data.get(f'{prefix}_hsn_code', ''),
                        'quantity': form_data.get(f'{prefix}_quantity', '1'),
                        'unit_value': form_data.get(f'{prefix}_unit_value', '0'),
                        'net_weight': form_data.get(f'{prefix}_net_weight', '0'),
                        'requester': form_data.get(f'{prefix}_attn', original_requester)
                    }
                    items_list.append(item_data)
                
                # Explicitly set the items as a list and count
                package_data['item_list'] = items_list
                package_data['items_count'] = len(items_list)
                print(f"DEBUG: Package {new_pkg_num} - type: {type(package_data)}")
                print(f"DEBUG: Package {new_pkg_num} - item_list type: {type(package_data['item_list'])}, length: {len(package_data['item_list'])}")
                print(f"DEBUG: Package {new_pkg_num} - keys: {list(package_data.keys())}")
                
                combined_packages.append(package_data)
            
            total_packages += current_packages
        
        print(f"DEBUG: Total combined packages: {total_packages}")
        print(f"DEBUG: Combined packages list length: {len(combined_packages)}")
        
    except Exception as e:
        print(f"DEBUG: Error processing packages: {e}")
        flash('Error processing shipment packages', 'error')
        return redirect(url_for('main.dashboard'))
    
    # Calculate total items
    total_items = sum(package.get('items_count', 0) for package in combined_packages)
    print(f"DEBUG: Total items: {total_items}")
    
    # Get all users for dropdown selection
    all_users = User.query.filter_by(is_active=True).order_by(User.first_name, User.last_name).all()
    
    # Prepare context for template
    context = {
        'shipments': shipments or [],  # Ensure it's always a list
        'combined_packages': combined_packages or [],  # Ensure it's always a list
        'combined_invoice': combined_invoice,
        'combined_number': combined_number,
        'total_packages': total_packages,
        'total_items': total_items,  # Pre-calculated total items
        'first_shipment': first_shipment,
        'shipment_ids': shipment_ids or [],  # Ensure it's always a list
        'all_users': all_users,  # Add users for dropdown
        'edit_mode': True,
        'shipment': shipments[0]
    }
    
    print(f"DEBUG: Template context prepared successfully")
    print(f"DEBUG: Shipments count: {len(context['shipments'])}")
    print(f"DEBUG: Combined packages count: {len(context['combined_packages'])}")
    
    return render_template('admin/combine_form.html', **context)

@main.route('/admin/finalize-combine', methods=['POST'])
@login_required
@admin_required
def admin_finalize_combine():
    """Finalize the combined shipment and generate document"""
    from flask import session
    
    try:
        shipment_ids = session.get('combine_shipment_ids', [])
        if not shipment_ids:
            flash('No shipments selected for combining', 'error')
            return redirect(url_for('main.dashboard'))
        
        # Get the form data from the combine form
        form_data = request.form.to_dict()
        
        # Process multi-value form fields (like checkboxes, multiple selects)
        for key, value in request.form.lists():
            if len(value) > 1:
                form_data[key] = value
        
        # Validate package limit for combined shipments (max 20 packages)
        total_packages = int(form_data.get('total_packages', 0))
        if total_packages > 20:
            flash('Error: Maximum 20 packages allowed for combined shipments.', 'error')
            return redirect(url_for('main.admin_combine_form'))
        
        # Get all shipments to combine
        shipments = Shipment.query.filter(Shipment.id.in_(shipment_ids)).all()
        first_shipment = shipments[0]
        
        # Get the combined invoice number from the form (already generated with proper format)
        combined_invoice = form_data.get('combined_invoice')
        
        # Extract serial number from the invoice number (it's the last part)
        serial_number = combined_invoice.split('/')[-1] if '/' in combined_invoice else '0001'
        
        # Create combined shipment record with additional fields
        combined_shipment = Shipment(
            invoice_number=combined_invoice,
            serial_number=serial_number,  # Use the extracted serial number
            shipment_type=first_shipment.shipment_type,
            status='Document_Generated',  # Combined shipments go straight to generated
            created_by=current_user.id,
            acknowledged_by=current_user.id,
            acknowledged_at=datetime.now(),
            requester_name=form_data.get('requester_name', 'Combined Shipment'),
            expedition_year=form_data.get('expedition_year', first_shipment.expedition_year),
            batch_number=first_shipment.batch_number,  # Use first shipment's batch number since it's removed from form
            destination_country=form_data.get('destination_country', first_shipment.destination_country),
            total_packages=int(form_data.get('total_packages', 0)),
            form_data=json.dumps(form_data),
            is_combined=True,
            combined_shipment_id=f"CMB{form_data.get('combined_number', '1')}",
            parent_shipment_ids=json.dumps(shipment_ids)
        )
        
        db.session.add(combined_shipment)
        db.session.flush()  # Flush to get the shipment ID before generating file reference
        
        # Generate file reference number for combined shipment
        combined_shipment.file_reference_number = generate_file_reference_number(combined_shipment, current_user)
        
        # Update original shipments status to 'Combined'
        for shipment in shipments:
            shipment.status = 'Combined'
        
        db.session.commit()
        
        # Clear session data
        session.pop('combine_shipment_ids', None)
        
        flash(f'Successfully combined {len(shipments)} shipments into {combined_invoice}! You can now generate documents from the dashboard.', 'success')
        
        # Redirect back to dashboard instead of immediately generating document
        return redirect(url_for('main.dashboard'))
        
    except Exception as e:
        db.session.rollback()
        flash(f'Error finalizing combined shipment: {str(e)}', 'error')
        return redirect(url_for('main.dashboard'))

@main.route('/admin/mark-delivered/<int:shipment_id>')
@login_required
@admin_required
def admin_mark_delivered(shipment_id):
    """Admin endpoint to mark a shipment as delivered to final destination"""
    shipment = Shipment.query.get_or_404(shipment_id)
    
    # Allow marking as delivered regardless of current status
    # Auto-acknowledge if not already acknowledged
    if not shipment.acknowledged_by:
        shipment.acknowledged_by = current_user.id
        shipment.acknowledged_at = datetime.now()
        # Generate file reference number if not already generated
        if not shipment.file_reference_number:
            shipment.file_reference_number = generate_file_reference_number(shipment, current_user)
    
    # Update shipment status to delivered
    shipment.status = 'Delivered'
    shipment.updated_at = datetime.now()
    db.session.commit()
    
    flash(f'Shipment {shipment.invoice_number} marked as delivered to final destination!', 'success')
    return redirect(url_for('main.dashboard'))

@main.route('/admin/delete-shipment/<int:shipment_id>')
@login_required
@admin_required
def admin_delete_shipment(shipment_id):
    """Admin endpoint to completely delete a shipment from the database"""
    shipment = Shipment.query.get_or_404(shipment_id)
    
    # Store shipment details for confirmation message
    invoice_number = shipment.invoice_number
    requester_name = shipment.requester_name
    
    try:
        # Delete the shipment completely from database
        db.session.delete(shipment)
        db.session.commit()
        
        flash(f'Shipment {invoice_number} (Requester: {requester_name}) has been permanently deleted from the database!', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'Error deleting shipment: {str(e)}', 'error')
    
    return redirect(url_for('main.dashboard'))

@main.route('/profile')
@login_required
def profile():
    """User profile page - view and edit profile information"""
    return render_template('auth/profile.html', user=current_user)

@main.route('/update-profile', methods=['POST'])
@login_required
def update_profile():
    """Update user profile information"""
    try:
        email = request.form.get('email')
        first_name = request.form.get('first_name')
        last_name = request.form.get('last_name')
        phone = request.form.get('phone')
        organization = request.form.get('organization')

        # Validate required fields
        if not all([email, first_name, last_name]):
            flash('Email, first name, and last name are required', 'error')
            return redirect(url_for('main.profile'))

        # Check if email is taken by another user
        if email != current_user.email:
            existing_user = User.query.filter_by(email=email).first()
            if existing_user:
                flash('Email address is already taken by another user', 'error')
                return redirect(url_for('main.profile'))

        # Update user information
        current_user.email = email
        current_user.first_name = first_name
        current_user.last_name = last_name
        current_user.phone = phone
        current_user.organization = organization

        db.session.commit()
        flash('Profile updated successfully!', 'success')
        
    except Exception as e:
        flash(f'Error updating profile: {str(e)}', 'error')
    
    return redirect(url_for('main.profile'))

@main.route('/change-password', methods=['POST'])
@login_required
def change_password():
    """Handle password change"""
    current_password = request.form.get('current_password')
    new_password = request.form.get('new_password')
    confirm_password = request.form.get('confirm_password')
    
    # Verify current password
    if not check_password_hash(current_user.password, current_password):
        flash('Current password is incorrect', 'error')
        return redirect(url_for('main.profile'))
    
    # Validate new password
    if new_password != confirm_password:
        flash('New passwords do not match', 'error')
        return redirect(url_for('main.profile'))
        
    if len(new_password) < 6:
        flash('Password must be at least 6 characters long', 'error')
        return redirect(url_for('main.profile'))
    
    # Update password
    current_user.password = generate_password_hash(new_password)
    db.session.commit()
    
    flash('Password updated successfully!', 'success')
    return redirect(url_for('main.profile'))

@main.route('/edit-shipment/<int:shipment_id>')
@login_required
def edit_shipment(shipment_id):
    """Edit shipment form - for both admin and non-admin users"""
    # Get the shipment
    shipment = Shipment.query.get_or_404(shipment_id)
    
    # Permission checks
    if current_user.is_admin():
        # Admins can edit any shipment (including combined ones)
        pass  # No restrictions for admins
    else:
        # Non-admin users can only edit their own shipments
        if shipment.created_by != current_user.id:
            flash('You can only edit your own shipments', 'error')
            return redirect(url_for('main.dashboard'))
        
        # Non-admin users cannot edit shipments in final states or combined shipments
        if shipment.status in ['Document_Generated', 'Delivered', 'Combined']:
            flash('This shipment cannot be edited as it has been processed or combined', 'error')
            return redirect(url_for('main.dashboard'))
        
        # Non-admin users cannot edit combined shipments
        if shipment.is_combined:
            flash('Combined shipments can only be edited by administrators', 'error')
            return redirect(url_for('main.dashboard'))
    
    # Parse form data
    form_data = json.loads(shipment.form_data) if shipment.form_data else {}
    
    # Ensure critical fields are available for form population
    if 'requester_user_id' not in form_data and shipment.created_by:
        form_data['requester_user_id'] = str(shipment.created_by)
    
    # Ensure expedition_month is available (extract from invoice if missing)
    if 'expedition_month' not in form_data and shipment.invoice_number:
        # Extract month from invoice number (format: NCPOR/ARC/YYYY/MMM/...)
        invoice_parts = shipment.invoice_number.split('/')
        if len(invoice_parts) >= 4:
            form_data['expedition_month'] = invoice_parts[3]  # MMM part
    
    # Ensure total_packages is available
    if 'total_packages' not in form_data and shipment.total_packages:
        form_data['total_packages'] = str(shipment.total_packages)
    
    # For combined shipments, use the combine form template
    if shipment.is_combined:
        # Prepare data structure similar to combine form
        combined_packages = []
        total_packages = int(form_data.get('total_packages', 0))
        
        for pkg_num in range(1, total_packages + 1):
            package_data = {
                'package_number': pkg_num,
                'type': get_package_type_display_name(form_data.get(f'package_{pkg_num}_type', ''), form_data, pkg_num),
                'length': form_data.get(f'package_{pkg_num}_length', ''),
                'width': form_data.get(f'package_{pkg_num}_width', ''),
                'height': form_data.get(f'package_{pkg_num}_height', ''),
                'item_list': [],
                'items_count': 0
            }
            
            # Get items for this package
            items_count = int(form_data.get(f'package_{pkg_num}_items_count', 0))
            items_list = []
            for item_num in range(1, items_count + 1):
                prefix = f'package_{pkg_num}_item_{item_num}'
                item_data = {
                    'description': form_data.get(f'{prefix}_description', ''),
                    'hsn_code': form_data.get(f'{prefix}_hsn_code', ''),
                    'quantity': form_data.get(f'{prefix}_quantity', '1'),
                    'unit_value': form_data.get(f'{prefix}_unit_value', '0'),
                    'net_weight': form_data.get(f'{prefix}_net_weight', '0'),
                    'requester': form_data.get(f'{prefix}_attn', 'Unknown')
                }
                items_list.append(item_data)
            
            package_data['item_list'] = items_list
            package_data['items_count'] = len(items_list)
            combined_packages.append(package_data)
        
        # Calculate total items
        total_items = sum(package.get('items_count', 0) for package in combined_packages)
        
        # Create context for combined shipment editing
        context = {
            'shipments': [shipment],  # Single combined shipment
            'combined_packages': combined_packages,
            'combined_invoice': shipment.invoice_number,
            'combined_number': shipment.combined_shipment_id.replace('CMB', '') if shipment.combined_shipment_id else '1',
            'total_packages': total_packages,
            'total_items': total_items,
            'first_shipment': shipment,
            'shipment_ids': [shipment.id],
            'edit_mode': True,
            'shipment': shipment
        }
        
        return render_template('admin/combine_form.html', **context)
    
    # For regular shipments, use normal templates
    if shipment.shipment_type == 'export':
        # Use different templates for admin vs regular users
        if current_user.is_admin():
            template = 'shipments/admin_export_shipment.html'
        else:
            template = 'shipments/export_shipment.html'
    else:
        # For other shipment types, use the standard templates
        template_map = {
            'import': 'shipments/import_shipment.html',
            'reimport': 'shipments/reimport_shipment.html',
            'cold': 'shipments/cold_shipment.html'
        }
        template = template_map.get(shipment.shipment_type, 'shipments/export_shipment.html')
    
    # Get all users for admin templates that need user selection
    all_users = User.query.filter_by(is_active=True).order_by(User.first_name, User.last_name).all()
    
    return render_template(template, 
                         edit_mode=True, 
                         shipment=shipment, 
                         form_data=form_data,
                         all_users=all_users)

@main.route('/update-shipment/<int:shipment_id>', methods=['POST'])
@login_required
def update_shipment(shipment_id):
    """Update shipment data - for both admin and non-admin users"""
    # Get the shipment
    shipment = Shipment.query.get_or_404(shipment_id)
    
    # Permission checks
    if current_user.is_admin():
        # Admins can update any shipment (including combined ones)
        pass  # No restrictions for admins
    else:
        # Non-admin users can only update their own shipments
        if shipment.created_by != current_user.id:
            flash('You can only update your own shipments', 'error')
            return redirect(url_for('main.dashboard'))
        
        # Non-admin users cannot update shipments in final states or combined shipments
        if shipment.status in ['Document_Generated', 'Delivered', 'Combined']:
            flash('This shipment cannot be updated as it has been processed or combined', 'error')
            return redirect(url_for('main.dashboard'))
        
        # Non-admin users cannot update combined shipments
        if shipment.is_combined:
            flash('Combined shipments can only be updated by administrators', 'error')
            return redirect(url_for('main.dashboard'))
    
    try:
        # Get all form data
        form_data = request.form.to_dict()
        
        # Convert multi-select values to lists
        for key in form_data:
            if key.endswith('[]'):
                form_data[key] = request.form.getlist(key)
        
        # Handle combined shipments differently
        if shipment.is_combined:
            # For combined shipments, don't regenerate invoice number - keep it the same
            # Just update the form data and other editable fields
            shipment.form_data = json.dumps(form_data)
            shipment.total_packages = int(form_data.get('total_packages', 0))
            shipment.updated_at = datetime.now()
            
            db.session.commit()
            
            flash(f'Combined shipment {shipment.invoice_number} updated successfully by Admin!', 'success')
            return redirect(url_for('main.dashboard'))
        
        # For regular shipments, handle normally
        # Extract basic shipment info
        requester_name = form_data.get('requester_name', '')
        expedition_year = form_data.get('expedition_year', '')
        expedition_month = form_data.get('expedition_month', '')
        destination_country = form_data.get('destination_country', '')
        total_packages = int(form_data.get('total_packages', 0))  # Use the form field directly
        
        # DO NOT regenerate invoice number when editing - preserve the original invoice number
        # The invoice number should remain the same as it was originally assigned
        
        # Update shipment (keeping original invoice number)
        shipment.requester_name = requester_name
        shipment.expedition_year = expedition_year
        shipment.expedition_month = form_data.get('expedition_month', '')
        shipment.destination_country = destination_country
        shipment.total_packages = total_packages
        
        # Update new shipping fields (only for admin users)
        if current_user.is_admin():
            shipment.country_of_final_destination = form_data.get('country_of_final_destination', '')
            shipment.mode_of_transport = form_data.get('mode_of_transport', '')
            shipment.port_of_loading = form_data.get('port_of_loading', '')
            shipment.port_of_discharge = form_data.get('port_of_discharge', '')
            shipment.final_destination = form_data.get('final_destination', '')
            shipment.country_of_origin = form_data.get('country_of_origin', 'India')
        
        shipment.form_data = json.dumps(form_data)
        
        # Status handling: Only reset status to 'Submitted' for non-admins
        if not current_user.is_admin():
            shipment.status = 'Submitted'  # Reset status for re-review for non-admins
            # Clear admin comments since shipment is being updated by non-admin
            shipment.admin_comment = None
            shipment.comment_by = None
            shipment.comment_at = None
        
        shipment.updated_at = datetime.now()
        
        db.session.commit()
        
        user_type = "Admin" if current_user.is_admin() else "User"
        flash(f'Shipment updated successfully by {user_type}! Invoice: {shipment.invoice_number} | Requester: {requester_name}, Year: {expedition_year}, Packages: {total_packages}', 'success')
        return redirect(url_for('main.dashboard'))
        
    except Exception as e:
        db.session.rollback()
        flash(f'Error updating shipment: {str(e)}', 'error')
        return redirect(url_for('main.dashboard'))

@main.route('/track-shipment/<int:shipment_id>')
@login_required
def track_shipment(shipment_id):
    """Track shipment page with visual progress indicator"""
    shipment = Shipment.query.get_or_404(shipment_id)
    
    # Permission check - users can only track their own shipments unless they're admin
    if not current_user.is_admin() and shipment.created_by != current_user.id:
        flash('You can only track your own shipments', 'error')
        return redirect(url_for('main.dashboard'))
    
    return render_template('shipments/tracking.html', shipment=shipment)

@main.route('/user/generate-document/<int:shipment_id>')
@main.route('/user/generate-document/<int:shipment_id>/<document_type>')
@login_required
def user_generate_document(shipment_id, document_type='invoice_packing'):
    """User endpoint to generate document for their own shipment"""
    shipment = Shipment.query.get_or_404(shipment_id)
    
    # Check if user owns this shipment
    if shipment.created_by != current_user.id:
        flash('You can only generate documents for your own shipments', 'error')
        return redirect(url_for('main.dashboard'))
    
    # Check if user has permission (non-admin users)
    if current_user.is_admin():
        flash('Admins use the admin dashboard to manage shipments', 'info')
        return redirect(url_for('main.dashboard'))
    
    # Validate document type
    if document_type not in ['invoice_packing', 'custom_docs']:
        document_type = 'invoice_packing'  # default fallback
    
    # Check if shipment can have documents generated
    if shipment.status not in ['Submitted', 'Acknowledged', 'Document_Generated']:
        flash('Cannot generate documents for this shipment status', 'error')
        return redirect(url_for('main.dashboard'))
    
    try:
        # Parse the stored form data
        form_data = json.loads(shipment.form_data)
        
        # Update status to document generated if it's not already
        if shipment.status in ['Submitted', 'Acknowledged']:
            original_status = shipment.status
            shipment.status = 'Document_Generated'
            if original_status == 'Submitted' and not shipment.acknowledged_by:
                # Auto-acknowledge if user is generating documents for submitted shipment
                shipment.acknowledged_by = current_user.id
                shipment.acknowledged_at = datetime.now()
            db.session.commit()
        
        # Generate and return the document with specified type
        return generate_shipment_document(shipment, form_data, document_type)
        
    except Exception as e:
        flash(f'Error generating document: {str(e)}', 'error')
        return redirect(url_for('main.dashboard'))

@main.route('/admin/update-status/<int:shipment_id>/<new_status>')
@login_required
@admin_required
def admin_update_status(shipment_id, new_status):
    """Admin endpoint to update shipment status to any of the 11 workflow steps"""
    shipment = Shipment.query.get_or_404(shipment_id)
    
    # Define valid status values
    valid_statuses = [
        'Acknowledged', 'Document_Generation', 'Quotation_Requested', 'Awaiting_Quotation_Approval'
    ]
    
    if new_status not in valid_statuses:
        flash('Invalid status provided', 'error')
        return redirect(url_for('main.dashboard'))
    
    # Auto-acknowledge if not already acknowledged (for any status beyond Submitted)
    if new_status != 'Submitted' and not shipment.acknowledged_by:
        shipment.acknowledged_by = current_user.id
        shipment.acknowledged_at = datetime.now()
        # Generate file reference number if not already generated
        if not shipment.file_reference_number:
            shipment.file_reference_number = generate_file_reference_number(shipment, current_user)
    
    # Update shipment status
    old_status = shipment.status
    shipment.status = new_status
    shipment.updated_at = datetime.now()
    db.session.commit()
    
    # Create user-friendly status names
    status_names = {
        'Acknowledged': 'Acknowledged',
        'Document_Generation': 'Generating Documents',
        'Quotation_Requested': 'Quotation Requested',
        'Awaiting_Quotation_Approval': 'Awaiting Quotation Approval'
    }
    
    # Check if this is an AJAX request
    if request.headers.get('Content-Type') == 'application/json' or request.headers.get('X-Requested-With') == 'XMLHttpRequest':
        return jsonify({
            'success': True,
            'message': f'Shipment {shipment.invoice_number} status updated from "{status_names.get(old_status, old_status)}" to "{status_names.get(new_status, new_status)}"',
            'new_status': new_status,
            'new_status_display': status_names.get(new_status, new_status)
        })
    else:
        flash(f'Shipment {shipment.invoice_number} status updated from "{status_names.get(old_status, old_status)}" to "{status_names.get(new_status, new_status)}"!', 'success')
        return redirect(url_for('main.dashboard'))

@main.route('/api/weather-proxy', methods=['GET'])
def weather_proxy():
    """Proxy endpoint for Yr.no weather API to handle CORS restrictions"""
    try:
        import requests
        from datetime import datetime, timedelta
        
        # Get coordinates from request args
        lat = request.args.get('lat', '78.9249')
        lon = request.args.get('lon', '11.9303')
        
        # Yr.no API URL
        yr_url = f'https://api.met.no/weatherapi/locationforecast/2.0/compact?lat={lat}&lon={lon}'
        
        # Set proper headers as required by Yr.no
        headers = {
            'User-Agent': 'COMPASS-Arctic-Research-Dashboard/1.0 (contact@compass-arctic.org)',
            'Accept': 'application/json'
        }
        
        print(f"Fetching weather from Yr.no: {yr_url}")
        
        # Make request to Yr.no
        response = requests.get(yr_url, headers=headers, timeout=10)
        
        print(f"Yr.no response status: {response.status_code}")
        
        if response.status_code == 200:
            weather_data = response.json()
            
            # Add CORS headers for browser access
            resp = jsonify({
                'success': True,
                'data': weather_data,
                'source': 'Yr.no (Norwegian Met Institute)',
                'timestamp': datetime.now().isoformat()
            })
            resp.headers['Access-Control-Allow-Origin'] = '*'
            resp.headers['Access-Control-Allow-Methods'] = 'GET'
            resp.headers['Access-Control-Allow-Headers'] = 'Content-Type'
            
            return resp
            
        else:
            print(f"Yr.no API error: {response.status_code} - {response.text}")
            return jsonify({
                'success': False,
                'error': f'Yr.no API returned status {response.status_code}',
                'status_code': response.status_code
            }), response.status_code
            
    except requests.exceptions.Timeout:
        print("Yr.no API timeout")
        return jsonify({
            'success': False,
            'error': 'Yr.no API timeout'
        }), 504
        
    except Exception as e:
        print(f"Weather proxy error: {e}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@main.route('/api/get-next-serial', methods=['GET'])
@login_required
def get_next_serial_preview():
    """
    Get the next serial number for preview without consuming it.
    Auto-optimizes counters if they're out of sync.
    """
    try:
        from compass.models import ShipmentSerialCounter, Shipment
        
        # Auto-optimize counter if needed
        shipment_count = Shipment.query.count()
        counter_record = ShipmentSerialCounter.query.first()
        
        if not counter_record:
            # Initialize counter based on actual shipments
            counter_record = ShipmentSerialCounter(counter=shipment_count)
            db.session.add(counter_record)
            db.session.commit()
            next_serial = f"{shipment_count + 1:04d}"
        elif counter_record.counter != shipment_count:
            # Counter is out of sync - optimize it
            counter_record.counter = shipment_count
            db.session.commit()
            next_serial = f"{shipment_count + 1:04d}"
        else:
            # Counter is in sync
            next_serial = f"{counter_record.counter + 1:04d}"
        
        return jsonify({
            'success': True,
            'next_serial': next_serial
        })
        
    except Exception as e:
        print(f"Error getting next serial: {e}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@main.route('/api/generate-invoice-preview', methods=['POST'])
@login_required
def generate_invoice_preview():
    """API endpoint to generate invoice number preview with unique ID and serial number"""
    try:
        data = request.get_json()
        
        if not data:
            return jsonify({'success': False, 'error': 'No data provided'})
        
        shipment_type = data.get('shipment_type', 'export')
        year = data.get('year', '')
        month = data.get('month', '')
        return_type = data.get('return_type', '')
        package_user_assignments = data.get('package_user_assignments', {})  # Dict of package_num: user_id
        selected_requester_id = data.get('selected_requester_id')  # Selected user ID from dropdown
        
        if not all([year, month, return_type]):
            return jsonify({'success': False, 'error': 'Missing required fields'})
        
        # Get next serial number - this is a preview, so we show the next available number
        from compass.models import ShipmentSerialCounter, User
        current_counter = ShipmentSerialCounter.query.first()
        next_serial = f"{(current_counter.counter + 1) if current_counter else 1:04d}"
        
        # Handle combined export for admin users
        if current_user.is_admin() and shipment_type == 'export' and package_user_assignments:
            # Extract unique user IDs in order of first package assignment
            unique_user_ids = []
            seen_users = set()
            
            # Sort by package number to maintain order of first assignment
            for package_num in sorted(package_user_assignments.keys(), key=int):
                user_id = package_user_assignments[package_num]
                if user_id and user_id not in seen_users:
                    user = User.query.get(user_id)
                    if user and user.unique_id:
                        unique_user_ids.append(user.unique_id)
                        seen_users.add(user_id)
            
            # If multiple users detected, create combined export
            if len(unique_user_ids) > 1:
                unique_ids_str = "/".join(unique_user_ids)
                invoice_number = f"NCPOR/ARC/{year}/{month}/EXP/{return_type}/CMB/{unique_ids_str}/{next_serial}"
                
                return jsonify({
                    'success': True,
                    'invoice_number': invoice_number,
                    'export_type': 'COMBINED',
                    'unique_ids': unique_user_ids,
                    'serial_number': next_serial,
                    'selected_users': len(unique_user_ids)
                })
            elif len(unique_user_ids) == 1:
                # Single user from package assignments
                user_unique_id = unique_user_ids[0]
                invoice_number = f"NCPOR/ARC/{year}/{month}/EXP/{return_type}/{user_unique_id}/{next_serial}"
                
                return jsonify({
                    'success': True,
                    'invoice_number': invoice_number,
                    'export_type': 'SINGLE',
                    'unique_id': user_unique_id,
                    'serial_number': next_serial
                })
        
        # Regular export or other shipment types (fallback)
        # Use selected requester's unique ID if provided, otherwise use current user
        if selected_requester_id:
            selected_user = User.query.get(selected_requester_id)
            user_unique_id = selected_user.unique_id if selected_user and selected_user.unique_id else 'XXXXXX'
            user_name = f"{selected_user.first_name} {selected_user.last_name}" if selected_user else 'Selected User'
        else:
            user_unique_id = current_user.unique_id if hasattr(current_user, 'unique_id') and current_user.unique_id else 'XXXXXX'
            user_name = f"{current_user.first_name} {current_user.last_name}"
        
        # Generate invoice number based on shipment type with serial number
        if shipment_type == 'export':
            invoice_number = f"NCPOR/ARC/{year}/{month}/EXP/{return_type}/{user_unique_id}/{next_serial}"
        elif shipment_type == 'import':
            # Use new NCPOR/ARC pattern for import shipments with admin ID and requester ID
            admin_id = current_user.unique_id if current_user.unique_id else 'ADMIN'
            invoice_number = f"NCPOR/ARC/{year}/{month}/SAM/RT/{admin_id}/{user_unique_id}/{next_serial}"
        elif shipment_type == 'reimport':
            invoice_number = f"NCPOR/REIMP/{year}/{month}/{return_type}/{user_unique_id}/{next_serial}"
        elif shipment_type == 'cold':
            invoice_number = f"NCPOR/COLD/{year}/{month}/{user_unique_id}/{next_serial}"
        else:
            invoice_number = f"NCPOR/UNKNOWN/{year}/{month}/{return_type}/{user_unique_id}/{next_serial}"
        
        return jsonify({
            'success': True,
            'invoice_number': invoice_number,
            'unique_id': user_unique_id,
            'serial_number': next_serial,
            'user_name': user_name
        })
        
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})