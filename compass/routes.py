from flask import Blueprint, render_template, request, send_file, flash, redirect, url_for, current_app
from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_ROW_HEIGHT
from docx.table import _Cell
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from datetime import datetime
from num2words import num2words
import os
import io
from flask_login import login_required, current_user

# Initialize the Blueprint for our main routes
main_bp = Blueprint('main', __name__)

def load_large_template():
    """
    Safely loads a large template file using buffered reading.
    
    Returns:
        Document: A python-docx Document object containing our template
    """
    try:
        # Get the absolute path to our template file
        current_dir = os.path.dirname(os.path.abspath(__file__))
        template_path = os.path.join(current_dir, '..', 'templates', 'export.docx')
        
        # Print debugging information to help troubleshoot path issues
        print(f"Current directory: {current_dir}")
        print(f"Attempting to load template from: {template_path}")
        print(f"File exists: {os.path.exists(template_path)}")
        
        if not os.path.exists(template_path):
            # If we can't find the file, let's see what's actually in the templates directory
            templates_dir = os.path.join(current_dir, '..', 'templates')
            if os.path.exists(templates_dir):
                print(f"Templates directory contents: {os.listdir(templates_dir)}")
            else:
                print("Templates directory not found!")
            raise FileNotFoundError(f"Template not found at {template_path}")
        
        # Create document from the template
        doc = Document(template_path)
        print("Successfully created Document object")
        return doc
            
    except Exception as e:
        print(f"Error in load_large_template: {str(e)}")
        raise

def handle_table_placement(doc, form_data):
    """
    Creates a cargo table with precise width specifications.
    """
    try:
        # Find the placeholder paragraph
        target_paragraph = None
        for paragraph in doc.paragraphs:
            if '[INVOICE_TABLE]' in paragraph.text:
                target_paragraph = paragraph
                break
        
        if not target_paragraph:
            raise ValueError("Could not find [INVOICE_TABLE] placeholder in template")

        # Create the table with the Table Grid style
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        table.autofit = False  # Prevent autofit from overriding our widths
        
        # Remove top border from the main table
        tblPr = table._element.xpath("./w:tblPr")[0]
        tblBorders = OxmlElement('w:tblBorders')
        
        # Set top border to none
        top = OxmlElement('w:top')
        top.set(qn('w:val'), 'none')
        tblBorders.append(top)
        
        # Set other borders (left, right, bottom, insideH, insideV) to single
        for border in ['left', 'right', 'bottom', 'insideH', 'insideV']:
            border_elem = OxmlElement(f'w:{border}')
            border_elem.set(qn('w:val'), 'single')
            border_elem.set(qn('w:sz'), '4')
            tblBorders.append(border_elem)
        
        tblPr.append(tblBorders)
        
        # Define column widths in centimeters
        columns_width = {
            0: 2.0,   # Marks & no
            1: 3.0,   # No & kind of packages
            2: 7.34,  # Description of goods
            3: 2.5,   # Quantity
            4: 2.25,  # Rate USD
            5: 2.25   # Amount USD
        }
        
        # Set table width and alignment
        table.width = Cm(19.34)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER  # Center align the entire table
        
        # Set up headers with proper formatting
        headers = ['Marks & no', 'No& kind of packages', 'Description of goods', 
                  'Quantity', 'Rate\nUSD', 'Amount\nUSD']
        
        for i, text in enumerate(headers):
            cell = table.cell(0, i)
            # Set column width
            cell.width = Cm(columns_width[i])
            
            paragraph = cell.paragraphs[0]
            paragraph.text = text
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Apply bold formatting to header text
            run = paragraph.runs[0]
                run.bold = True
            run.font.size = Pt(11)  # Set font size
            
            # Set vertical alignment to center
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            tcVAlign = OxmlElement('w:vAlign')
            tcVAlign.set(qn('w:val'), "center")
            tcPr.append(tcVAlign)

        # Insert the table after the placeholder and remove the placeholder
        table._element = target_paragraph._element.addnext(table._element)
        target_paragraph._element.getparent().remove(target_paragraph._element)
        
        return table
        
    except Exception as e:
        print(f"Error in handle_table_placement: {str(e)}")
        raise

def populate_table_data(table, form_data):
    """
    Populates the table with form data while maintaining precise formatting.
    """
    try:
        # Define column widths in centimeters
        columns_width = {
            0: 2.0,   # Marks & no
            1: 3.0,   # No & kind of packages
            2: 7.34,  # Description of goods
            3: 2.5,   # Quantity
            4: 2.25,  # Rate USD
            5: 2.25   # Amount USD
        }

        # Calculate the total width of columns 3-5 for the nested table
        nested_table_width = sum(columns_width[i] for i in range(3, 6))

        total_amount = 0
        total_packages = int(form_data.get('total_packages', 0))

        # Process each package and its items
        for package_num in range(1, total_packages + 1):
            items_count = int(form_data.get(f'package_{package_num}_items_count', 0))
            
            for item_num in range(1, items_count + 1):
                # Add new row for each item
                row = table.add_row()
                row_cells = row.cells
                prefix = f'package_{package_num}_item_{item_num}'
                
                # Apply column widths to the new row
                for idx, width in columns_width.items():
                    cell = row_cells[idx]
                    cell.width = Cm(width)
                
                # Serial number
                row_cells[0].text = f"{package_num}."
                row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                # Box details
                package_type = form_data.get(f'package_{package_num}_type', '')
                row_cells[1].text = f"Box-{package_num}\n({package_type})"
                row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                # Description of goods with HSN Code
                description = form_data.get(f'{prefix}_description', '')
                hsn_code = form_data.get(f'{prefix}_hsn_code', '')
                row_cells[2].text = f"{description}\nHSN: {hsn_code}"
                row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                # Quantity
                quantity = int(form_data.get(f'{prefix}_quantity', 0))
                row_cells[3].text = f"{quantity} SET"
                row_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                # Rate USD
                unit_value = float(form_data.get(f'{prefix}_unit_value', 0))
                row_cells[4].text = f"{unit_value:.0f}"
                row_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                
                # Amount USD
                amount = unit_value * quantity
                row_cells[5].text = f"{amount:.0f}"
                row_cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                
                total_amount += amount

        # Convert number to words and capitalize first letter of each word, and "And"
        amount_words = num2words(int(total_amount), to='cardinal').split()
        capitalized_words = []
        for word in amount_words:
            if word.lower() == 'and':
                capitalized_words.append('And')
            else:
                capitalized_words.append(word.capitalize())
        amount_in_words = " ".join(capitalized_words) + " Only"

        # Create the declarations text with proper formatting
        declarations_text = (
            f"Amount Chargeable (in words): USD {amount_in_words}\n\n"
            "THE VALUE DECLARED ABOVE IS FOR CUSTOMS PURPOSE ONLY.\n\n"
            "GOODS ARE FOR RESEARCH AND DEVELOPMENT.\n\n"
            "Declaration:\n"
            "We declare that the invoice shows the actual price of goods\n"
            "Described and that all particulars are true and correct."
        )

        # Add final row for declarations
        final_row = table.add_row()
        final_cells = final_row.cells
        
        # Apply column widths
        for idx, width in columns_width.items():
            cell = final_cells[idx]
            cell.width = Cm(width)

        # Set up declarations section (columns 0-2)
        final_cells[0].text = declarations_text
        for i in range(1, 3):
            final_cells[0].merge(final_cells[i])
        final_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Merge the cells for the right side first
        final_cells[3].merge(final_cells[4])
        final_cells[3].merge(final_cells[5])

        # Get the exact width of the merged cell in twips (1/20th of a point)
        merged_cell_width = int(nested_table_width * 1440)  # Convert cm to twips
        
        # Create a 2x1 table for the right side (TOTAL and Signature)
        right_table = Document().add_table(rows=2, cols=2)
        right_table.style = 'Table Grid'
        right_table.autofit = False

        # Set table properties to match container cell exactly and center align the table
        tblPr = right_table._element.xpath("./w:tblPr")[0]
        
        # Remove all outer borders from the table
        tblBorders = OxmlElement('w:tblBorders')
        for border in ['top', 'bottom', 'left', 'right']:
            border_elem = OxmlElement(f'w:{border}')
            border_elem.set(qn('w:val'), 'none')
            tblBorders.append(border_elem)
        tblPr.append(tblBorders)
        
        # Convert centimeters to twips (1cm = 567 twips)
        total_width = int(7.0 * 567)  # Total table width of 7cm
        tblW = parse_xml(f'<w:tblW xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:w="{total_width}" w:type="dxa"/>')
        jc = parse_xml('<w:jc xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="center"/>')
        tblPr.append(tblW)
        tblPr.append(jc)
        
        # Set exact row heights
        total_row = right_table.rows[0]
        signature_row = right_table.rows[1]
        
        # Set TOTAL row height to 0.8cm
        total_row.height_rule = WD_ROW_HEIGHT.EXACTLY
        total_row.height = Cm(0.8)
        
        # Set signature row height to (3.64 - 0.8) = 2.84cm
        signature_row.height_rule = WD_ROW_HEIGHT.EXACTLY
        signature_row.height = Cm(2.84)

        # Set up TOTAL row with exact column widths
        total_label_cell = right_table.rows[0].cells[0]
        total_amount_cell = right_table.rows[0].cells[1]
        
        # Convert cell widths to twips
        total_label_width = int(2.5 * 567)  # 2.5cm for TOTAL label
        total_amount_width = int(4.5 * 567)  # 4.5cm for amount
        
        # Set exact widths for both cells
        total_label_cell._tc.get_or_add_tcPr().append(
            parse_xml(f'<w:tcW xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:w="{total_label_width}" w:type="dxa"/>')
        )
        total_amount_cell._tc.get_or_add_tcPr().append(
            parse_xml(f'<w:tcW xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:w="{total_amount_width}" w:type="dxa"/>')
        )
        
        # Format TOTAL label
        total_label_paragraph = total_label_cell.paragraphs[0]
        total_label_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = total_label_paragraph.add_run("TOTAL")
        run.font.bold = True
        
        # Format amount
        amount_paragraph = total_amount_cell.paragraphs[0]
        amount_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = amount_paragraph.add_run(f"{total_amount:.0f}")
        run.font.bold = True

        # Set up Signature row
        sig_cell = right_table.rows[1].cells[0]
        sig_cell.merge(right_table.rows[1].cells[1])
        
        # Set width for merged signature cell (7cm)
        sig_cell_width = int(7.0 * 567)  # 7cm
        sig_cell._tc.get_or_add_tcPr().append(
            parse_xml(f'<w:tcW xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:w="{sig_cell_width}" w:type="dxa"/>')
        )
        
        sig_paragraph = sig_cell.paragraphs[0]
        sig_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        sig_paragraph.add_run("Signature & date")
        
        # Add extra paragraphs for signature space
        for _ in range(2):
            sig_cell.add_paragraph()

        # Remove any existing content from container cell
        final_cells[3]._element.clear_content()
        
        # Convert the nested table to XML and insert it
        nested_table_xml = right_table._element
        final_cells[3]._tc.append(nested_table_xml)

        # Remove all borders from both tables
        # Main table borders
        for row in table.rows:
            for cell in row.cells:
                if hasattr(cell._tc, 'tcPr'):
                    cell._tc.get_or_add_tcPr().append(
                        OxmlElement('w:tcBorders')
                    )

        # Nested table borders
        for row in right_table.rows:
            for cell in row.cells:
                if hasattr(cell._tc, 'tcPr'):
                    cell._tc.get_or_add_tcPr().append(
                        OxmlElement('w:tcBorders')
                    )

        # Add bottom border under TOTAL row using Word's built-in properties
        def set_bottom_border(cell):
            """Set bottom border for a cell using Word's built-in properties"""
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '4')
            tcBorders.append(bottom)
            tcPr.append(tcBorders)

        # Apply bottom border to both TOTAL cells
        set_bottom_border(total_label_cell)
        set_bottom_border(total_amount_cell)

        return total_amount

    except Exception as e:
        print(f"Error in populate_table_data: {str(e)}")
        raise

def handle_pl_table_placement(doc, form_data):
    """
    Creates a packing list table with precise width specifications.
    """
    try:
        # Find the placeholder paragraph
        target_paragraph = None
        for paragraph in doc.paragraphs:
            if '[PL_TABLE]' in paragraph.text:
                target_paragraph = paragraph
                break
        
        if not target_paragraph:
            raise ValueError("Could not find [PL_TABLE] placeholder in template")

        # Create the table with the Table Grid style
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        table.autofit = False  # Prevent autofit from overriding our widths
        
        # Remove top border from the main table
        tblPr = table._element.xpath("./w:tblPr")[0]
        tblBorders = OxmlElement('w:tblBorders')
        
        # Set top border to none
        top = OxmlElement('w:top')
        top.set(qn('w:val'), 'none')
        tblBorders.append(top)
        
        # Set other borders (left, right, bottom, insideH, insideV) to single
        for border in ['left', 'right', 'bottom', 'insideH', 'insideV']:
            border_elem = OxmlElement(f'w:{border}')
            border_elem.set(qn('w:val'), 'single')
            border_elem.set(qn('w:sz'), '4')
            tblBorders.append(border_elem)
        
        tblPr.append(tblBorders)
        
        # Define column widths in centimeters for PL table
        columns_width = {
            0: 2.0,   # Marks & no
            1: 3.0,   # No & kind of packages
            2: 7.34,  # Description of goods
            3: 2.0,   # Quantity
            4: 3.0,   # Dimensions
            5: 2.0    # Weight
        }
        
        # Set table width and alignment
        table.width = Cm(19.34)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER  # Center align the entire table
        
        # Set up headers with proper formatting
        headers = ['Marks & no', 'No& kind of packages', 'Description of goods', 
                  'Quantity', 'Dimension\n(L x B x H)\n(cm)', 'Weight\n(kg)']
        
        for i, text in enumerate(headers):
            cell = table.cell(0, i)
            # Set column width
            cell.width = Cm(columns_width[i])
            
            paragraph = cell.paragraphs[0]
            paragraph.text = text
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Apply bold formatting to header text
            run = paragraph.runs[0]
            run.bold = True
            run.font.size = Pt(11)  # Set font size
            
            # Set vertical alignment to center
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            tcVAlign = OxmlElement('w:vAlign')
            tcVAlign.set(qn('w:val'), "center")
            tcPr.append(tcVAlign)

        # Insert the table after the placeholder and remove the placeholder
        table._element = target_paragraph._element.addnext(table._element)
        target_paragraph._element.getparent().remove(target_paragraph._element)
        
        return table
        
    except Exception as e:
        print(f"Error in handle_pl_table_placement: {str(e)}")
        raise

def populate_pl_table_data(table, form_data):
    """
    Populates the packing list table with form data while maintaining precise formatting.
    """
    try:
        # Define column widths in centimeters
        columns_width = {
            0: 2.0,   # Marks & no
            1: 3.0,   # No & kind of packages
            2: 7.34,  # Description of goods
            3: 2.0,   # Quantity
            4: 3.0,   # Dimensions
            5: 2.0    # Weight
        }

        total_packages = int(form_data.get('total_packages', 0))

        # Process each package and its items
        for package_num in range(1, total_packages + 1):
            items_count = int(form_data.get(f'package_{package_num}_items_count', 0))
            
            for item_num in range(1, items_count + 1):
                # Add new row for each item
                row = table.add_row()
                row_cells = row.cells
                prefix = f'package_{package_num}_item_{item_num}'
                
                # Apply column widths to the new row
                for idx, width in columns_width.items():
                    cell = row_cells[idx]
                    cell.width = Cm(width)
                
                # Serial number
                row_cells[0].text = f"{package_num}."
                row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                # Box details
                package_type = form_data.get(f'package_{package_num}_type', '')
                row_cells[1].text = f"Box-{package_num}\n({package_type})"
                row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                # Description of goods with HSN Code
                description = form_data.get(f'{prefix}_description', '')
                hsn_code = form_data.get(f'{prefix}_hsn_code', '')
                row_cells[2].text = f"{description}\nHSN: {hsn_code}"
                row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                # Quantity
                quantity = int(form_data.get(f'{prefix}_quantity', 0))
                row_cells[3].text = f"{quantity}"
                row_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Dimensions
                length = form_data.get(f'package_{package_num}_length', '')
                width = form_data.get(f'package_{package_num}_width', '')
                height = form_data.get(f'package_{package_num}_height', '')
                row_cells[4].text = f"{length} x {width} x {height}"
                row_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Weight
                weight = float(form_data.get(f'{prefix}_net_weight', 0))
                row_cells[5].text = f"{weight:.2f}"
                row_cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Add final row for declaration and signature
        final_row = table.add_row()
        final_cells = final_row.cells
        
        # Apply column widths
        for idx, width in columns_width.items():
            cell = final_cells[idx]
            cell.width = Cm(width)

        # Merge cells for declaration (left side)
        final_cells[0].merge(final_cells[1])
        final_cells[0].merge(final_cells[2])
        
        # Merge cells for signature (right side)
        final_cells[3].merge(final_cells[4])
        final_cells[3].merge(final_cells[5])
        
        # Add declaration text on the left
        declaration_text = (
            "THE ABOVE MENTIONED GOODS ARE FOR RESEARCH AND DEVELOPMENT.\n\n"
            "Declaration:\n"
            "We declare that all particulars given above are true and correct."
        )
        final_cells[0].text = declaration_text
        final_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Add signature text on the right
        final_cells[3].text = "Signature & date"
        final_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Set row height for signature space
        final_row.height = Cm(3.0)
        final_row.height_rule = WD_ROW_HEIGHT.EXACTLY

        return None

    except Exception as e:
        print(f"Error in populate_pl_table_data: {str(e)}")
        raise

@main_bp.route('/submit-shipment', methods=['POST'])
def submit_shipment():
    """
    Handles form submission and document generation using both docxtpl and python-docx.
    """
    try:
        # Get form data
        form_data = request.form.to_dict()

        # Get the absolute path to our template file
        current_dir = os.path.dirname(os.path.abspath(__file__))
        template_path = os.path.join(current_dir, '..', 'templates', 'export.docx')

        # Load the template with docxtpl for initial replacements
        tpl = DocxTemplate(template_path)

        # Helper function to get first name from full name
        def get_first_name(full_name):
            if not full_name:
                return ''
            parts = full_name.split()
            # If we have more than one word (likely a title is present)
            if len(parts) > 1:
                return parts[1].upper()  # Take second word as first name
            return parts[0].upper()  # If only one word, use that

        # Prepare context for template rendering
        context = {
            # Invoice Number Generation
            'invoice_no': f"NCPOR/ARC/{form_data.get('expedition_year', '')}/{form_data.get('return_type', '')}/{form_data.get('batch_number', '')}/{form_data.get('requester_name', '').split()[1].upper() if len(form_data.get('requester_name', '').split()) > 1 else form_data.get('requester_name', '').split()[0].upper() if form_data.get('requester_name', '') else 'FIRSTNAME'}",
            
            # Current Date
            'invoice_date': datetime.now().strftime('%d-%m-%Y'),
            
            # Exporter Details
            'exporter_name': "National Center for Polar and Ocean Research (NCPOR) [erstwhite NCAOR]",
            'exporter_ministry': "Ministry of Earth Sciences (Govt. of India)",
            'exporter_address': "Headland Sada, Vasco da Gama, Goa - 403804",
            'exporter_country': "India",
            'exporter_phone': "+91 832 2525501",
            'exporter_gst': "30AACFN4991PlZN",
            'exporter_iec': "1196000125",
            
            # Consignee Details
            'consignee_name': "Himadri - Indian Arctic Research Station C/O Kingsbay AS",
            'consignee_address': "N-9173 Ny-Alesund, Norway",
            'consignee_contact': "Kingsbay AS, Longyearbyen",
            'consignee_phone': "+47 79 027200",
            
            # Shipment Details
            'destination_country': form_data.get('destination_country', 'NORWAY'),
            'airport_of_loading': form_data.get('airport_loading', 'MUMBAI'),
            'requester_name': form_data.get('requester_name', ''),
            'expedition_year': form_data.get('expedition_year', ''),
            'batch_number': form_data.get('batch_number', ''),
            'return_type': form_data.get('return_type', ''),
            
            # Additional Details
            'purpose': "RESEARCH AND DEVELOPMENT",
            'declaration': "We declare that the invoice shows the actual price of goods described and that all particulars are true and correct."
        }

        # Render the template with docxtpl
        tpl.render(context)

        # Save the docxtpl output to a temporary file
        temp_path = os.path.join(current_dir, '..', 'templates', 'temp_export.docx')
        tpl.save(temp_path)

        # Load the temporary file with python-docx for table manipulation
        doc = Document(temp_path)

        # Handle table placement and population
        invoice_table = handle_table_placement(doc, form_data)
        total_amount = populate_table_data(invoice_table, form_data)

        # Handle packing list table
        pl_table = handle_pl_table_placement(doc, form_data)
        populate_pl_table_data(pl_table, form_data)

        # Compute amount in words
        amount_in_words = f"USD {num2words(int(total_amount))} Only"

        # Update amount in words
        for paragraph in doc.paragraphs:
            if '[AMOUNT_IN_WORDS]' in paragraph.text:
                paragraph.text = paragraph.text.replace('[AMOUNT_IN_WORDS]', amount_in_words)

        # Save to a buffer
        docx_buffer = io.BytesIO()
        doc.save(docx_buffer)
        docx_buffer.seek(0)

        # Remove the temporary file
        os.remove(temp_path)

        # Generate a unique filename
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"export_invoice_{timestamp}.docx"

        # Return the generated document
        return send_file(
            docx_buffer,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=filename
        )

    except Exception as e:
        error_message = f"Error generating document: {str(e)}"
        print(error_message)
        flash(error_message, 'error')
        return redirect(url_for('main.index'))

@main_bp.route('/submit-import-shipment', methods=['POST'])
def submit_import_shipment():
    """
    Handles import form submission and document generation.
    """
    try:
        # Get form data
        form_data = request.form.to_dict()

        # Get the absolute path to our template file
        current_dir = os.path.dirname(os.path.abspath(__file__))
        template_path = os.path.join(current_dir, '..', 'templates', 'export.docx')  # Reusing export template for now

        # Load the template with docxtpl for initial replacements
        tpl = DocxTemplate(template_path)

        # Prepare context for template rendering
        context = {
            # Invoice Number Generation
            'invoice_no': f"NCPOR/IMP/{form_data.get('expedition_year', '')}/{form_data.get('import_purpose', 'RESEARCH')}/{form_data.get('batch_number', '')}/{form_data.get('requester_name', '').split()[1].upper() if len(form_data.get('requester_name', '').split()) > 1 else form_data.get('requester_name', '').split()[0].upper() if form_data.get('requester_name', '') else 'FIRSTNAME'}",
            
            # Current Date
            'invoice_date': datetime.now().strftime('%d-%m-%Y'),
            
            # Exporter Details (Himadri Station for imports)
            'exporter_name': "Station Leader",
            'exporter_org': "HIMADRI-Indian Arctic Research Station",
            'exporter_address': "Ny-Alesund, c/o Kings Bay AS",
            'exporter_location': "N-9173, Ny-Alesund, Norway",
            'exporter_phone': "+47 79 02 72 00",
            
            # Consignee Details (NCPOR for imports)
            'consignee_name': "National Center for Polar and Ocean Research (NCPOR) [erstwhite NCAOR]",
            'consignee_ministry': "Ministry of Earth Sciences (Govt. of India)",
            'consignee_address': "Headland Sada, Vasco da Gama, Goa - 403804",
            'consignee_location': "India",
            'consignee_phone': "+91 832 2525501",
            'consignee_gst': "30AACFN4991PlZN",
            'consignee_iec': "1196000125",
            
            # Consigner Details (Himadri Station for imports)
            'consigner_name': "Himadri - Indian Arctic Research Station C/O Kingsbay AS",
            'consigner_address': "N-9173 Ny-Alesund, Norway",
            'consigner_contact': "Kingsbay AS, Longyearbyen",
            'consigner_phone': "+47 79 027200",
            
            # Import Details
            'destination_country': 'INDIA',
            'airport_of_loading': form_data.get('import_mode', 'AIR').upper(),
            'requester_name': form_data.get('requester_name', ''),
            'expedition_year': form_data.get('expedition_year', ''),
            'batch_number': form_data.get('batch_number', ''),
            'import_purpose': form_data.get('import_purpose', 'RESEARCH'),
            
            # Additional Details
            'purpose': "RESEARCH AND DEVELOPMENT",
            'declaration': "We declare that the invoice shows the actual price of goods described and that all particulars are true and correct."
        }

        # Render the template with docxtpl
        tpl.render(context)

        # Save the docxtpl output to a temporary file
        temp_path = os.path.join(current_dir, '..', 'templates', 'temp_import.docx')
        tpl.save(temp_path)

        # Load the temporary file with python-docx for table manipulation
        doc = Document(temp_path)

        # Handle table placement and population
        table = handle_table_placement(doc, form_data)
        total_amount = populate_table_data(table, form_data)

        # Compute amount in words
        amount_in_words = f"USD {num2words(int(total_amount))} Only"

        # Update amount in words
        for paragraph in doc.paragraphs:
            if '[AMOUNT_IN_WORDS]' in paragraph.text:
                paragraph.text = paragraph.text.replace('[AMOUNT_IN_WORDS]', amount_in_words)

        # Save to a buffer
        docx_buffer = io.BytesIO()
        doc.save(docx_buffer)
        docx_buffer.seek(0)

        # Remove the temporary file
        os.remove(temp_path)

        # Generate a unique filename
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"import_invoice_{timestamp}.docx"

        # Return the generated document
        return send_file(
            docx_buffer,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=filename
        )

    except Exception as e:
        error_message = f"Error generating document: {str(e)}"
        print(error_message)
        flash(error_message, 'error')
        return redirect(url_for('main.index'))

@main_bp.route('/')
def index():
    """Renders the home page"""
    return render_template('index.html', now=datetime.now())

@main_bp.route('/shipment-type')
def shipment_type_selection():
    """Renders the shipment type selection page"""
    return render_template('shipments/type_selection.html')

@main_bp.route('/new-shipment')
def new_shipment():
    """
    Handles the shipment form display based on type.
    Validates the shipment type and shows the appropriate form.
    """
    shipment_type = request.args.get('type', 'cold')
    templates = {
        'export': 'shipments/export_shipment.html',
        'import': 'shipments/import_shipment.html',
        'reimport': 'shipments/reimport_shipment.html',
        'cold': 'shipments/cold_shipment.html'
    }
    template = templates.get(shipment_type)
    if not template:
        flash('Invalid shipment type selected', 'error')
        return redirect(url_for('main.shipment_type_selection'))
    return render_template(template)

@main_bp.route('/dashboard')
def dashboard():
    """
    Displays the shipment tracking dashboard with filters and status.
    In a real application, this would fetch data from a database.
    """
    # Mock data for demonstration purposes
    shipment_data = {
        'active_shipments': 8,
        'completed_shipments': 12,
        'delayed_shipments': 2,
        'total_value': '78,342'
    }
    
    return render_template('shipments/dashboard.html', **shipment_data)